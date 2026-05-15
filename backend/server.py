"""Mind Matters — Personal Operating System backend.

FastAPI + MongoDB. Single-user v1 with multi-user-ready schema (every doc has user_id).
All routes are prefixed with /api. Times are stored as ISO strings. _id is excluded.
"""
from pathlib import Path
from dotenv import load_dotenv

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / ".env")  # must run BEFORE tg/docs_gen imports

from fastapi import FastAPI, APIRouter, HTTPException, Depends, UploadFile, File, Form, Header, Query
from fastapi.responses import JSONResponse
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
from pydantic import BaseModel, Field, ConfigDict
from typing import List, Optional, Dict, Any, Literal
from datetime import datetime, timezone, timedelta, date
import os
import io
import asyncio
import base64
import time
import uuid
import secrets
import logging
import json
import jwt as pyjwt
import httpx

from emergentintegrations.llm.chat import LlmChat, UserMessage
import bcrypt

from docs_gen import render_by_template_id, BUILTIN_TEMPLATES, render_simple_statement
from tg import tg_send, tg_send_document, tg_poll_loop, reminder_loop

MONGO_URL = os.environ["MONGO_URL"]
DB_NAME = os.environ["DB_NAME"]
JWT_SECRET = os.environ.get("JWT_SECRET", "mind-matters-dev-secret")
EMERGENT_LLM_KEY = os.environ.get("EMERGENT_LLM_KEY", "")

client = AsyncIOMotorClient(MONGO_URL)
db = client[DB_NAME]

app = FastAPI(title="Mind Matters API", version="1.0.0")
api = APIRouter(prefix="/api")

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("mind-matters")


# ───────────────────────────── helpers ─────────────────────────────
def now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def today_key() -> str:
    return datetime.now(timezone.utc).date().isoformat()


def to_iso(v):
    if isinstance(v, datetime):
        return v.isoformat()
    return v


def new_id() -> str:
    return str(uuid.uuid4())


def make_token(user_id: str) -> str:
    payload = {
        "user_id": user_id,
        "iat": int(datetime.now(timezone.utc).timestamp()),
        "exp": int((datetime.now(timezone.utc) + timedelta(days=30)).timestamp()),
    }
    return pyjwt.encode(payload, JWT_SECRET, algorithm="HS256")


async def get_current_user(authorization: Optional[str] = Header(None)) -> dict:
    """Extract user from Authorization: Bearer <token>. Auto-creates demo user if none."""
    token = None
    if authorization and authorization.lower().startswith("bearer "):
        token = authorization.split(" ", 1)[1].strip()
    if not token:
        raise HTTPException(status_code=401, detail="Not authenticated")
    try:
        payload = pyjwt.decode(token, JWT_SECRET, algorithms=["HS256"])
    except Exception:
        raise HTTPException(status_code=401, detail="Invalid token")
    user = await db.users.find_one({"id": payload["user_id"]}, {"_id": 0})
    if not user:
        raise HTTPException(status_code=401, detail="User not found")
    return user


# ───────────────────────────── models ─────────────────────────────
class User(BaseModel):
    id: str
    email: str
    first_name: str
    last_name: str = ""
    picture: str = ""
    created_at: str


class DemoLoginReq(BaseModel):
    first_name: Optional[str] = "Friend"


class SignupReq(BaseModel):
    first_name: str
    email: str
    password: str


class LoginReq(BaseModel):
    email: str
    password: str


def _hash_password(password: str) -> str:
    return bcrypt.hashpw(password.encode("utf-8"), bcrypt.gensalt()).decode("utf-8")


def _verify_password(password: str, hashed: str) -> bool:
    try:
        return bcrypt.checkpw(password.encode("utf-8"), hashed.encode("utf-8"))
    except Exception:
        return False


class TokenResp(BaseModel):
    token: str
    user: User


# Task
class TaskIn(BaseModel):
    date: Optional[str] = None
    name: str = ""  # responsible person / "to"
    task: str
    details: str = ""
    status: str = "Pending"  # free text — common values: Pending, Completed, Done (legacy), Follow-Up, or any custom user value
    group: str = ""  # custom group, user-defined (replaces block1..4)
    section: str = ""  # optional second-level grouping (Todoist-style sub-header)
    parent_id: Optional[str] = None  # if set, this is a subtask nested under another task
    flagged: bool = False  # priority flag — flagged rows sort to top
    project_id: Optional[str] = None  # v2.17 — project this row belongs to
    # Lightweight inline attachments — each entry: {name, mime, size, data_url}.
    # Capped at ~4MB total per task in the upload endpoint.
    attachments: List[Dict[str, Any]] = []


class Task(TaskIn):
    id: str
    sr_no: int
    order_index: int = 0
    user_id: str
    created_at: str
    updated_at: str


# Routine
class RoutineIn(BaseModel):
    group: str = ""  # custom group name (e.g. "Morning", "4 Hours Focus")
    name: str = ""  # custom sub-label (e.g. person or area of life)
    activity: str
    details: str = ""
    frequency: str = "Daily"  # free-form now (Daily/Weekly/Custom…)
    priority: Literal["Low", "Medium", "High"] = "Medium"
    status: Literal["Active", "Paused"] = "Active"
    section: str = ""  # optional second-level grouping (Todoist-style sub-header)
    parent_id: Optional[str] = None  # nested subtask support (max depth 3, enforced client-side)
    flagged: bool = False
    project_id: Optional[str] = None  # v2.17 — project this row belongs to
    attachments: List[Dict[str, Any]] = []


class Routine(RoutineIn):
    id: str
    sr_no: int = 0
    order_index: int = 0
    user_id: str
    created_at: str
    updated_at: str


class RoutineLogIn(BaseModel):
    routine_id: str
    date: Optional[str] = None  # YYYY-MM-DD
    done: bool = True


class RoutineLog(BaseModel):
    id: str
    user_id: str
    routine_id: str
    date: str
    done: bool
    created_at: str


# Loan
class LoanIn(BaseModel):
    date: Optional[str] = None
    name: str
    amount: float
    interest: float = 0.0  # value (rate% if interest_type='percent', flat ₹ if 'fixed')
    interest_type: Literal["percent", "fixed"] = "percent"
    reason: str = ""
    status: Literal["Given", "Taken", "Pending", "Closed"] = "Given"
    repayment_date: Optional[str] = None


class Loan(LoanIn):
    id: str
    sr_no: int
    user_id: str
    created_at: str
    updated_at: str


# Transactions (unified Cash Flow — replaces Loans/Investments/Insurance)
class TransactionIn(BaseModel):
    date: Optional[str] = None
    amount: float
    name: str = ""  # legacy / kept for back-compat
    vendor: str = ""  # v2.2 — preferred display label
    details: str = ""
    remarks: str = ""  # legacy
    mode: str = "Cash"  # v2.2 — preferred for "mode of payment"
    head: str = "Uncategorized"
    category: str = "expense"  # was Literal — now free text so users can add custom
    group: str = ""
    section: str = ""  # second-level grouping
    company: str = ""
    expense_head: str = "Uncategorized"
    direction: Literal["in", "out"] = "out"
    account: str = "Personal"
    notes: str = ""
    # Loan-style fields (used when category=liability or asset for loans given)
    interest_rate: Optional[float] = None  # annual % rate
    interest_type: Literal["percent", "fixed"] = "percent"
    repayment_date: Optional[str] = None  # YYYY-MM-DD
    emi: Optional[float] = None  # monthly EMI / installment ₹
    currency: str = "INR"  # ISO code per row; summary cards convert to user's default
    parent_id: Optional[str] = None  # for splitting bills into line items
    flagged: bool = False
    project_id: Optional[str] = None  # v2.17 — project this row belongs to
    attachments: List[Dict[str, Any]] = []
    source: Literal["manual", "upload", "telegram"] = "manual"


class Transaction(TransactionIn):
    id: str
    sr_no: int = 0
    order_index: int = 0
    user_id: str
    created_at: str
    updated_at: str


# Investments
class InvestmentIn(BaseModel):
    kind: Literal["investment", "insurance"] = "investment"
    type: str = "Equity"  # free-form: Equity / FD / Insurance / MF / custom
    provider: str
    amount_invested: float
    start_date: Optional[str] = None
    maturity_date: Optional[str] = None
    rate_or_value: Optional[str] = None
    current_value: Optional[float] = None
    insured_for: Optional[str] = None  # only when kind='insurance'
    notes: str = ""


class Investment(InvestmentIn):
    id: str
    sr_no: int
    user_id: str
    created_at: str
    updated_at: str


# Notes
class NoteIn(BaseModel):
    title: str = ""
    body: str = ""
    tags: List[str] = []
    pinned: bool = False
    flagged: bool = False
    project_id: Optional[str] = None  # v2.17 — project this row belongs to
    attachments: List[Dict[str, Any]] = []


class Note(NoteIn):
    id: str
    user_id: str
    created_at: str
    updated_at: str


# Affirmation
class AffirmationIn(BaseModel):
    date: Optional[str] = None
    text: str = ""
    personal_fixed: Optional[str] = None  # user's fixed personal affirmation


class Affirmation(BaseModel):
    id: str
    user_id: str
    date: str
    text: str = ""
    personal_fixed: str = ""
    updated_at: str


# ───────────────────────────── auth ─────────────────────────────
@api.post("/auth/demo-login", response_model=TokenResp)
async def demo_login(body: DemoLoginReq):
    """Instant auth for v1. Creates or reuses the 'You' user."""
    first_name = (body.first_name or "Friend").strip() or "Friend"
    user = await db.users.find_one({"email": "you@mindmatters.local"}, {"_id": 0})
    if not user:
        user = {
            "id": new_id(),
            "email": "you@mindmatters.local",
            "first_name": first_name,
            "last_name": "",
            "picture": "",
            "created_at": now_iso(),
        }
        await db.users.insert_one(dict(user))
    else:
        if user.get("first_name") != first_name:
            await db.users.update_one({"id": user["id"]}, {"$set": {"first_name": first_name}})
            user["first_name"] = first_name
    token = make_token(user["id"])
    return {"token": token, "user": User(**user)}


@api.get("/auth/me", response_model=User)
async def auth_me(user=Depends(get_current_user)):
    return User(**user)


# ───────────────────────────── auth: email/password ─────────────────────────────
@api.post("/auth/signup", response_model=TokenResp)
async def signup(body: SignupReq):
    email = body.email.strip().lower()
    if not email or "@" not in email:
        raise HTTPException(400, "Invalid email")
    if len(body.password) < 6:
        raise HTTPException(400, "Password must be at least 6 characters")
    first_name = body.first_name.strip() or "Friend"

    existing = await db.users.find_one({"email": email}, {"_id": 0})
    if existing and existing.get("password_hash"):
        raise HTTPException(400, "Email already registered — use login instead")

    # If this is the first real signup and a demo-login user exists with no password,
    # attach the signup to that same user_id so historical data stays linked.
    demo = await db.users.find_one(
        {"email": "you@mindmatters.local", "password_hash": {"$in": [None, ""]}},
        {"_id": 0},
    )
    if demo and not existing:
        await db.users.update_one(
            {"id": demo["id"]},
            {"$set": {
                "email": email,
                "first_name": first_name,
                "password_hash": _hash_password(body.password),
            }},
        )
        user = await db.users.find_one({"id": demo["id"]}, {"_id": 0})
    else:
        user = {
            "id": new_id(),
            "email": email,
            "first_name": first_name,
            "last_name": "",
            "picture": "",
            "password_hash": _hash_password(body.password),
            "created_at": now_iso(),
        }
        await db.users.insert_one(dict(user))
    token = make_token(user["id"])
    user_resp = {k: v for k, v in user.items() if k != "password_hash"}
    return {"token": token, "user": User(**user_resp)}


@api.post("/auth/login", response_model=TokenResp)
async def login(body: LoginReq):
    email = body.email.strip().lower()
    user = await db.users.find_one({"email": email}, {"_id": 0})
    if not user or not user.get("password_hash"):
        raise HTTPException(401, "Invalid email or password")
    if not _verify_password(body.password, user["password_hash"]):
        raise HTTPException(401, "Invalid email or password")
    token = make_token(user["id"])
    user_resp = {k: v for k, v in user.items() if k != "password_hash"}
    return {"token": token, "user": User(**user_resp)}


# ───────────────────────── Password reset ─────────────────────────
class ForgotReq(BaseModel):
    email: str


class ResetReq(BaseModel):
    email: str
    code: str
    new_password: str = Field(min_length=6, max_length=128)


@api.post("/auth/forgot")
async def auth_forgot(body: ForgotReq):
    """Generate a 6-digit reset code, store it (hashed), and deliver it
    out-of-band via Telegram (if the user has linked their bot). The code is
    NEVER returned in the response body — that would defeat the security model.
    Response shape: {ok, delivered_via, expires_at}.
    """
    email = body.email.strip().lower()
    user = await db.users.find_one({"email": email}, {"_id": 0, "id": 1})
    # Always succeed so we don't leak whether an email is registered.
    if not user:
        return {"ok": True, "delivered_via": "none"}
    import secrets
    code = f"{secrets.randbelow(10**6):06d}"
    expires_iso = (datetime.now(timezone.utc) + timedelta(minutes=30)).isoformat()
    await db.password_resets.update_one(
        {"user_id": user["id"]},
        {"$set": {
            "user_id": user["id"],
            "email": email,
            "code_hash": _hash_password(code),
            "expires_at": expires_iso,
            "consumed": False,
            "created_at": now_iso(),
        }},
        upsert=True,
    )
    # If Telegram is linked for this user, also DM the code to them.
    delivered_via = "screen"
    try:
        tg_link = await db.tg_links.find_one({"user_id": user["id"]}, {"_id": 0})
        if tg_link and tg_link.get("chat_id"):
            from tg import tg_send
            await tg_send(
                tg_link["chat_id"],
                f"🔐 Mind Matters password reset code: *{code}*\nExpires in 30 min.",
                parse_mode="Markdown",
            )
            delivered_via = "telegram+screen"
    except Exception as e:
        logger.warning(f"forgot-password TG send failed: {e}")
    return {"ok": True, "delivered_via": delivered_via, "expires_at": expires_iso}


@api.post("/auth/reset", response_model=TokenResp)
async def auth_reset(body: ResetReq):
    email = body.email.strip().lower()
    user = await db.users.find_one({"email": email}, {"_id": 0})
    if not user:
        raise HTTPException(400, "Invalid code")
    rec = await db.password_resets.find_one({"user_id": user["id"], "consumed": False}, {"_id": 0})
    if not rec:
        raise HTTPException(400, "Invalid or expired code")
    if rec.get("expires_at", "") < datetime.now(timezone.utc).isoformat():
        raise HTTPException(400, "Code expired — request a new one")
    if not _verify_password(body.code.strip(), rec["code_hash"]):
        raise HTTPException(400, "Invalid code")
    # Apply new password
    await db.users.update_one(
        {"id": user["id"]},
        {"$set": {"password_hash": _hash_password(body.new_password), "updated_at": now_iso()}},
    )
    await db.password_resets.update_one(
        {"user_id": user["id"]}, {"$set": {"consumed": True, "consumed_at": now_iso()}}
    )
    fresh = await db.users.find_one({"id": user["id"]}, {"_id": 0, "password_hash": 0})
    token = make_token(user["id"])
    return {"token": token, "user": User(**fresh)}


# ───────────────────────────── tasks ─────────────────────────────
@api.get("/tasks", response_model=List[Task])
async def list_tasks(
    status: Optional[str] = None,
    name: Optional[str] = None,
    date_from: Optional[str] = None,
    date_to: Optional[str] = None,
    project_id: Optional[str] = None,
    user=Depends(get_current_user),
):
    q: Dict[str, Any] = {"user_id": user["id"]}
    if status:
        q["status"] = status
    if name:
        q["name"] = name
    if date_from or date_to:
        q["date"] = {}
        if date_from:
            q["date"]["$gte"] = date_from
        if date_to:
            q["date"]["$lte"] = date_to
    if project_id:
        q["project_id"] = project_id
    docs = await db.tasks.find(q, {"_id": 0}).sort("sr_no", 1).to_list(5000)
    return [Task(**d) for d in docs]


async def _next_sr(collection: str, user_id: str) -> int:
    doc = await db[collection].find_one(
        {"user_id": user_id}, {"_id": 0, "sr_no": 1}, sort=[("sr_no", -1)]
    )
    return (doc["sr_no"] + 1) if doc else 1


async def _compact_sr(collection: str, user_id: str):
    """Re-number sr_no contiguously after a delete."""
    docs = await db[collection].find(
        {"user_id": user_id}, {"_id": 0, "id": 1}
    ).sort("sr_no", 1).to_list(20000)
    for i, d in enumerate(docs, 1):
        await db[collection].update_one({"id": d["id"]}, {"$set": {"sr_no": i}})


async def _resequence_sr(collection: str, user_id: str, target_id: str, new_sr: int):
    """Move target_id's sr_no to new_sr and shift the rest to keep numbering contiguous."""
    docs = await db[collection].find(
        {"user_id": user_id}, {"_id": 0, "id": 1, "sr_no": 1}
    ).sort("sr_no", 1).to_list(20000)
    # Pull target out
    others = [d for d in docs if d.get("id") != target_id]
    # Clamp new_sr within [1, len(others) + 1]
    new_sr = max(1, min(new_sr, len(others) + 1))
    insert_at = new_sr - 1
    others.insert(insert_at, {"id": target_id})
    for i, d in enumerate(others, 1):
        await db[collection].update_one(
            {"id": d["id"], "user_id": user_id}, {"$set": {"sr_no": i}}
        )


@api.post("/tasks", response_model=Task)
async def create_task(body: TaskIn, user=Depends(get_current_user)):
    sr = await _next_sr("tasks", user["id"])
    doc = {
        **body.model_dump(),
        "id": new_id(),
        "sr_no": sr,
        "user_id": user["id"],
        "date": body.date or today_key(),
        "created_at": now_iso(),
        "updated_at": now_iso(),
    }
    if not doc.get("project_id"):
        doc["project_id"] = await _ensure_default_project(user["id"])
    await _assert_project_write(user["id"], doc["project_id"])
    await db.tasks.insert_one(dict(doc))
    return Task(**doc)


@api.post("/tasks/bulk", response_model=List[Task])
async def bulk_create_tasks(body: List[TaskIn], user=Depends(get_current_user)):
    out = []
    sr = await _next_sr("tasks", user["id"])
    default_pid = await _ensure_default_project(user["id"])
    for b in body:
        doc = {
            **b.model_dump(),
            "id": new_id(),
            "sr_no": sr,
            "user_id": user["id"],
            "date": b.date or today_key(),
            "created_at": now_iso(),
            "updated_at": now_iso(),
        }
        if not doc.get("project_id"):
            doc["project_id"] = default_pid
        sr += 1
        await db.tasks.insert_one(dict(doc))
        out.append(Task(**doc))
    return out


@api.patch("/tasks/{task_id}", response_model=Task)
async def update_task(task_id: str, body: Dict[str, Any], user=Depends(get_current_user)):
    body = {k: v for k, v in body.items()
            if k in {"date", "name", "task", "details", "status", "group", "section",
                     "order_index", "sr_no", "parent_id", "attachments", "flagged",
                     "project_id"}}
    # Status compatibility: accept "Completed" alongside legacy "Done"
    if body.get("status") == "Completed":
        body["status"] = "Completed"  # keep as-is
    body["updated_at"] = now_iso()
    # If sr_no is being explicitly edited, re-sequence siblings around the new value.
    if "sr_no" in body:
        try:
            new_sr = max(1, int(body["sr_no"]))
        except Exception:
            body.pop("sr_no", None)
            new_sr = None
        if new_sr is not None:
            await _resequence_sr("tasks", user["id"], task_id, new_sr)
            body["sr_no"] = new_sr
    res = await db.tasks.find_one_and_update(
        {"id": task_id, "user_id": user["id"]},
        {"$set": body},
        projection={"_id": 0},
        return_document=True,
    )
    if not res:
        raise HTTPException(404, "Task not found")
    return Task(**res)


@api.delete("/tasks/{task_id}")
async def delete_task(task_id: str, user=Depends(get_current_user)):
    # Cascade: also delete subtasks pointing to this task as parent
    await db.tasks.delete_many({"user_id": user["id"], "parent_id": task_id})
    r = await db.tasks.delete_one({"id": task_id, "user_id": user["id"]})
    if r.deleted_count == 0:
        raise HTTPException(404, "Task not found")
    await _compact_sr("tasks", user["id"])
    return {"ok": True}


@api.post("/tasks/{task_id}/attachments")
async def upload_task_attachment(
    task_id: str,
    file: UploadFile = File(...),
    user=Depends(get_current_user),
):
    """Attach a small file (<4MB) inline to a task as base64 data_url."""
    t = await db.tasks.find_one({"id": task_id, "user_id": user["id"]}, {"_id": 0})
    if not t:
        raise HTTPException(404, "Task not found")
    content = await file.read()
    if len(content) > 4 * 1024 * 1024:
        raise HTTPException(413, "Attachment too large (max 4MB)")
    mime = file.content_type or "application/octet-stream"
    data_url = f"data:{mime};base64,{base64.b64encode(content).decode()}"
    entry = {
        "id": new_id()[:10],
        "name": file.filename or "file",
        "mime": mime,
        "size": len(content),
        "data_url": data_url,
    }
    current = t.get("attachments") or []
    current.append(entry)
    # Cap total inline attachments at ~8MB to keep document size reasonable
    total = sum(a.get("size", 0) for a in current)
    if total > 8 * 1024 * 1024:
        raise HTTPException(413, "Total attachments would exceed 8MB")
    await db.tasks.update_one(
        {"id": task_id, "user_id": user["id"]},
        {"$set": {"attachments": current, "updated_at": now_iso()}},
    )
    return entry


@api.delete("/tasks/{task_id}/attachments/{att_id}")
async def delete_task_attachment(task_id: str, att_id: str, user=Depends(get_current_user)):
    t = await db.tasks.find_one({"id": task_id, "user_id": user["id"]}, {"_id": 0})
    if not t:
        raise HTTPException(404, "Task not found")
    remaining = [a for a in (t.get("attachments") or []) if a.get("id") != att_id]
    await db.tasks.update_one(
        {"id": task_id, "user_id": user["id"]},
        {"$set": {"attachments": remaining, "updated_at": now_iso()}},
    )
    return {"ok": True}


# ───────────────────────────── routines ─────────────────────────────
@api.get("/routines", response_model=List[Routine])
async def list_routines(project_id: Optional[str] = None, user=Depends(get_current_user)):
    q: Dict[str, Any] = {"user_id": user["id"]}
    if project_id:
        q["project_id"] = project_id
    docs = await db.routines.find(q, {"_id": 0}).sort("sr_no", 1).to_list(1000)
    # Backfill sr_no for any old docs missing it (one-time, ordered by created_at).
    if any((d.get("sr_no") or 0) <= 0 for d in docs):
        ordered = sorted(docs, key=lambda d: d.get("created_at", ""))
        for i, d in enumerate(ordered, 1):
            if (d.get("sr_no") or 0) != i:
                await db.routines.update_one({"id": d["id"]}, {"$set": {"sr_no": i}})
                d["sr_no"] = i
        docs = sorted(ordered, key=lambda d: d.get("sr_no", 0))
    return [Routine(**d) for d in docs]


@api.post("/routines", response_model=Routine)
async def create_routine(body: RoutineIn, user=Depends(get_current_user)):
    sr = await _next_sr("routines", user["id"])
    doc = {
        **body.model_dump(),
        "id": new_id(),
        "sr_no": sr,
        "user_id": user["id"],
        "created_at": now_iso(),
        "updated_at": now_iso(),
    }
    if not doc.get("project_id"):
        doc["project_id"] = await _ensure_default_project(user["id"])
    await _assert_project_write(user["id"], doc["project_id"])
    await db.routines.insert_one(dict(doc))
    return Routine(**doc)


@api.post("/routines/bulk", response_model=List[Routine])
async def bulk_routines(body: List[RoutineIn], user=Depends(get_current_user)):
    out = []
    sr = await _next_sr("routines", user["id"])
    default_pid = await _ensure_default_project(user["id"])
    for b in body:
        doc = {
            **b.model_dump(),
            "id": new_id(),
            "sr_no": sr,
            "user_id": user["id"],
            "created_at": now_iso(),
            "updated_at": now_iso(),
        }
        if not doc.get("project_id"):
            doc["project_id"] = default_pid
        sr += 1
        await db.routines.insert_one(dict(doc))
        out.append(Routine(**doc))
    return out


@api.patch("/routines/{rid}", response_model=Routine)
async def update_routine(rid: str, body: Dict[str, Any], user=Depends(get_current_user)):
    body = {k: v for k, v in body.items()
            if k in {"group", "name", "activity", "details", "frequency", "priority", "status",
                     "section", "order_index", "sr_no", "parent_id", "flagged", "attachments",
                     "project_id"}}
    body["updated_at"] = now_iso()
    if "sr_no" in body:
        try:
            new_sr = max(1, int(body["sr_no"]))
        except Exception:
            body.pop("sr_no", None)
            new_sr = None
        if new_sr is not None:
            await _resequence_sr("routines", user["id"], rid, new_sr)
            body["sr_no"] = new_sr
    res = await db.routines.find_one_and_update(
        {"id": rid, "user_id": user["id"]}, {"$set": body}, projection={"_id": 0}, return_document=True
    )
    if not res:
        raise HTTPException(404, "Routine not found")
    return Routine(**res)


@api.delete("/routines/{rid}")
async def delete_routine(rid: str, user=Depends(get_current_user)):
    await db.routines.delete_one({"id": rid, "user_id": user["id"]})
    await db.routine_logs.delete_many({"routine_id": rid, "user_id": user["id"]})
    await _compact_sr("routines", user["id"])
    return {"ok": True}


@api.post("/routine-logs", response_model=RoutineLog)
async def toggle_routine_log(body: RoutineLogIn, user=Depends(get_current_user)):
    d = body.date or today_key()
    existing = await db.routine_logs.find_one(
        {"user_id": user["id"], "routine_id": body.routine_id, "date": d}, {"_id": 0}
    )
    if existing:
        await db.routine_logs.update_one(
            {"id": existing["id"]}, {"$set": {"done": body.done}}
        )
        existing["done"] = body.done
        return RoutineLog(**existing)
    doc = {
        "id": new_id(),
        "user_id": user["id"],
        "routine_id": body.routine_id,
        "date": d,
        "done": body.done,
        "created_at": now_iso(),
    }
    await db.routine_logs.insert_one(dict(doc))
    return RoutineLog(**doc)


@api.get("/routine-logs")
async def list_routine_logs(
    date_from: Optional[str] = None, date_to: Optional[str] = None, user=Depends(get_current_user)
):
    q: Dict[str, Any] = {"user_id": user["id"]}
    if date_from or date_to:
        q["date"] = {}
        if date_from:
            q["date"]["$gte"] = date_from
        if date_to:
            q["date"]["$lte"] = date_to
    docs = await db.routine_logs.find(q, {"_id": 0}).to_list(10000)
    return docs


def _calc_streak(logs_by_date: Dict[str, bool]) -> int:
    """Count consecutive daily completions ending today or yesterday."""
    today = datetime.now(timezone.utc).date()
    streak = 0
    d = today
    # allow yesterday too
    if not logs_by_date.get(d.isoformat()):
        d = d - timedelta(days=1)
    while logs_by_date.get(d.isoformat()):
        streak += 1
        d = d - timedelta(days=1)
    return streak


@api.get("/routines/summary")
async def routines_summary(user=Depends(get_current_user)):
    routines = await db.routines.find({"user_id": user["id"]}, {"_id": 0}).to_list(1000)
    logs = await db.routine_logs.find({"user_id": user["id"]}, {"_id": 0}).to_list(50000)
    today = today_key()
    per_routine = {}
    for r in routines:
        r_logs = {l["date"]: l["done"] for l in logs if l["routine_id"] == r["id"]}
        per_routine[r["id"]] = {
            "routine": r,
            "done_today": bool(r_logs.get(today)),
            "streak": _calc_streak(r_logs),
        }
    total = len(routines)
    done_today = sum(1 for v in per_routine.values() if v["done_today"])
    percent_today = round((done_today / total) * 100) if total else 0
    # per category %
    cats: Dict[str, Dict[str, int]] = {}
    for r in routines:
        c = r.get("group") or r.get("time_block") or r.get("category") or "General"
        cats.setdefault(c, {"total": 0, "done": 0})
        cats[c]["total"] += 1
        if per_routine[r["id"]]["done_today"]:
            cats[c]["done"] += 1
    category_percent = {
        c: round((v["done"] / v["total"]) * 100) if v["total"] else 0 for c, v in cats.items()
    }
    return {
        "total": total,
        "done_today": done_today,
        "percent_today": percent_today,
        "category_percent": category_percent,
        "per_routine": per_routine,
    }


# ───────────────────────────── generic reorder + groups ─────────────────────────────
class ReorderReq(BaseModel):
    ids: List[str]  # desired order


_REORDER_COLLECTIONS = {"tasks", "routines", "transactions"}


@api.post("/{resource}/reorder")
async def reorder_items(resource: str, body: ReorderReq, user=Depends(get_current_user)):
    if resource not in _REORDER_COLLECTIONS:
        raise HTTPException(400, "Unsupported resource")
    coll = db[resource]
    # Update both order_index AND sr_no so visible Sr column reflects new order.
    for idx, _id in enumerate(body.ids):
        await coll.update_one(
            {"id": _id, "user_id": user["id"]},
            {"$set": {"order_index": idx, "sr_no": idx + 1, "updated_at": now_iso()}},
        )
    return {"ok": True, "count": len(body.ids)}


@api.get("/groups/{resource}")
async def list_groups(resource: str, user=Depends(get_current_user)):
    """Distinct custom group names used by this user for a given resource."""
    if resource not in _REORDER_COLLECTIONS:
        raise HTTPException(400, "Unsupported resource")
    coll = db[resource]
    names = await coll.distinct("group", {"user_id": user["id"]})
    return {"groups": sorted([n for n in names if n])}


# ───────────────────────────── loans ─────────────────────────────
def _accrued_interest(amount: float, rate: float, start_iso: Optional[str], itype: str = "percent") -> float:
    if not start_iso or not rate:
        return 0.0
    if itype == "fixed":
        # flat interest amount (treat as already-accrued total)
        return round(float(rate), 2)
    try:
        start = datetime.fromisoformat(start_iso).date() if "T" in start_iso else date.fromisoformat(start_iso)
    except Exception:
        return 0.0
    days = (datetime.now(timezone.utc).date() - start).days
    if days <= 0:
        return 0.0
    return round(amount * (rate / 100.0) * (days / 365.0), 2)


@api.get("/loans")
async def list_loans(user=Depends(get_current_user)):
    docs = await db.loans.find({"user_id": user["id"]}, {"_id": 0}).sort("sr_no", 1).to_list(5000)
    for d in docs:
        d["accrued_interest"] = _accrued_interest(
            d["amount"], d.get("interest", 0), d.get("date"), d.get("interest_type", "percent")
        )
    return docs


@api.post("/loans", response_model=Loan)
async def create_loan(body: LoanIn, user=Depends(get_current_user)):
    sr = await _next_sr("loans", user["id"])
    doc = {
        **body.model_dump(),
        "id": new_id(),
        "sr_no": sr,
        "user_id": user["id"],
        "date": body.date or today_key(),
        "created_at": now_iso(),
        "updated_at": now_iso(),
    }
    await db.loans.insert_one(dict(doc))
    return Loan(**doc)


@api.patch("/loans/{lid}", response_model=Loan)
async def update_loan(lid: str, body: Dict[str, Any], user=Depends(get_current_user)):
    allowed = {"date", "name", "amount", "interest", "interest_type", "reason", "status", "repayment_date"}
    body = {k: v for k, v in body.items() if k in allowed}
    body["updated_at"] = now_iso()
    res = await db.loans.find_one_and_update(
        {"id": lid, "user_id": user["id"]}, {"$set": body}, projection={"_id": 0}, return_document=True
    )
    if not res:
        raise HTTPException(404, "Loan not found")
    return Loan(**res)


@api.delete("/loans/{lid}")
async def delete_loan(lid: str, user=Depends(get_current_user)):
    await db.loans.delete_one({"id": lid, "user_id": user["id"]})
    await _compact_sr("loans", user["id"])
    return {"ok": True}


@api.get("/loans/summary")
async def loans_summary(user=Depends(get_current_user)):
    loans = await db.loans.find({"user_id": user["id"]}, {"_id": 0}).to_list(5000)
    given = sum(l["amount"] for l in loans if l["status"] == "Given")
    taken = sum(l["amount"] for l in loans if l["status"] == "Taken")
    pending = [l for l in loans if l["status"] in ("Given", "Taken", "Pending")]
    total_interest = sum(
        _accrued_interest(l["amount"], l.get("interest", 0), l.get("date"), l.get("interest_type", "percent"))
        for l in pending
    )
    overdue = 0
    today = datetime.now(timezone.utc).date()
    for l in pending:
        rep = l.get("repayment_date")
        if rep:
            try:
                rd = date.fromisoformat(rep)
                if rd < today:
                    overdue += 1
            except Exception:
                pass
    return {
        "total_given": given,
        "total_taken": taken,
        "net_exposure": given - taken,
        "total_interest_accrued": round(total_interest, 2),
        "overdue_count": overdue,
        "count": len(loans),
    }


# ───────────────────────────── transactions (cash flow) ─────────────────────────────
@api.get("/transactions")
async def list_transactions(
    month: Optional[str] = None,  # YYYY-MM
    expense_head: Optional[str] = None,
    project_id: Optional[str] = None,
    user=Depends(get_current_user),
):
    q: Dict[str, Any] = {"user_id": user["id"]}
    if month:
        q["date"] = {"$regex": f"^{month}"}
    if expense_head:
        q["expense_head"] = expense_head
    if project_id:
        q["project_id"] = project_id
    docs = await db.transactions.find(q, {"_id": 0}).sort("sr_no", 1).to_list(10000)
    # Backfill sr_no for any legacy docs missing it.
    if any((d.get("sr_no") or 0) <= 0 for d in docs):
        ordered = sorted(docs, key=lambda d: d.get("created_at", ""))
        for i, d in enumerate(ordered, 1):
            if (d.get("sr_no") or 0) != i:
                await db.transactions.update_one({"id": d["id"]}, {"$set": {"sr_no": i}})
                d["sr_no"] = i
        docs = sorted(ordered, key=lambda d: d.get("sr_no", 0))
    return docs


@api.post("/transactions", response_model=Transaction)
async def create_transaction(body: TransactionIn, user=Depends(get_current_user)):
    sr = await _next_sr("transactions", user["id"])
    doc = {
        **body.model_dump(),
        "id": new_id(),
        "sr_no": sr,
        "user_id": user["id"],
        "date": body.date or today_key(),
        "created_at": now_iso(),
        "updated_at": now_iso(),
    }
    if not doc.get("project_id"):
        doc["project_id"] = await _ensure_default_project(user["id"])
    await _assert_project_write(user["id"], doc["project_id"])
    await db.transactions.insert_one(dict(doc))
    return Transaction(**doc)


@api.patch("/transactions/{tid}", response_model=Transaction)
async def update_transaction(tid: str, body: Dict[str, Any], user=Depends(get_current_user)):
    allowed = {"date", "amount", "mode", "company", "expense_head", "direction",
               "account", "notes", "name", "vendor", "details", "remarks", "head",
               "category", "group", "section", "order_index", "sr_no",
               "interest_rate", "interest_type", "repayment_date", "emi",
               "currency", "parent_id", "flagged", "attachments", "project_id"}
    body = {k: v for k, v in body.items() if k in allowed}
    body["updated_at"] = now_iso()
    if "sr_no" in body:
        try:
            new_sr = max(1, int(body["sr_no"]))
        except Exception:
            body.pop("sr_no", None)
            new_sr = None
        if new_sr is not None:
            await _resequence_sr("transactions", user["id"], tid, new_sr)
            body["sr_no"] = new_sr
    res = await db.transactions.find_one_and_update(
        {"id": tid, "user_id": user["id"]}, {"$set": body}, projection={"_id": 0}, return_document=True
    )
    if not res:
        raise HTTPException(404, "Transaction not found")
    return Transaction(**res)


@api.delete("/transactions/{tid}")
async def delete_transaction(tid: str, user=Depends(get_current_user)):
    await db.transactions.delete_one({"id": tid, "user_id": user["id"]})
    await _compact_sr("transactions", user["id"])
    return {"ok": True}


def _parse_excel(content: bytes) -> List[Dict[str, Any]]:
    import pandas as pd
    df = pd.read_excel(io.BytesIO(content))
    df.columns = [str(c).strip().lower() for c in df.columns]
    rows = df.to_dict(orient="records")
    return rows


def _parse_pdf(content: bytes) -> str:
    import pdfplumber
    text = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages[:10]:
            t = page.extract_text() or ""
            text.append(t)
    return "\n".join(text)


async def _ai_categorize(rows: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """Send parsed rows to Gemini and normalize into TransactionIn-like dicts."""
    if not EMERGENT_LLM_KEY:
        return []
    import json
    prompt = (
        "You are an AI accountant. Given a list of raw expense/income rows, "
        "return a JSON array of objects with keys: date (YYYY-MM-DD), amount (number), "
        "mode, company, expense_head (short category like Food, Travel, Utilities, Rent, Salary, Misc), "
        "direction ('in' for income, 'out' for expense), notes. "
        "Return ONLY a JSON array, no prose.\n\nRows:\n" + json.dumps(rows[:50], default=str)
    )
    chat = LlmChat(
        api_key=EMERGENT_LLM_KEY,
        session_id=f"parse-{new_id()}",
        system_message="You parse raw financial records into clean JSON.",
    ).with_model("gemini", "gemini-3-flash-preview")
    resp = await chat.send_message(UserMessage(text=prompt))
    text = resp.strip()
    # strip code fences if any
    if text.startswith("```"):
        text = text.strip("`")
        if text.startswith("json"):
            text = text[4:]
    try:
        parsed = json.loads(text)
        if isinstance(parsed, list):
            return parsed
    except Exception as e:
        logger.error(f"AI parse JSON failed: {e}; raw={text[:400]}")
    return []


@api.post("/transactions/upload")
async def upload_transactions(
    file: UploadFile = File(...),
    account: str = Form("Personal"),
    skip_duplicates: bool = Form(True),
    user=Depends(get_current_user),
):
    content = await file.read()
    name = (file.filename or "").lower()
    raw_rows: List[Dict[str, Any]] = []
    if name.endswith((".xlsx", ".xls")):
        try:
            raw_rows = _parse_excel(content)
        except Exception as e:
            raise HTTPException(400, f"Failed to parse Excel: {e}")
    elif name.endswith(".pdf"):
        text = _parse_pdf(content)
        raw_rows = [{"line": line} for line in text.split("\n") if line.strip()][:100]
    elif name.endswith(".csv"):
        import csv
        f = io.StringIO(content.decode("utf-8", errors="ignore"))
        raw_rows = list(csv.DictReader(f))
    else:
        raise HTTPException(400, "Unsupported file type. Upload .xlsx, .xls, .csv or .pdf")

    categorized = await _ai_categorize(raw_rows)
    inserted = []
    pending = []  # duplicates - to be confirmed by user
    next_sr = await _next_sr("transactions", user["id"])
    for row in categorized:
        try:
            amount = float(row.get("amount", 0) or 0)
        except Exception:
            continue
        if amount == 0:
            continue
        d = str(row.get("date") or today_key())[:10]
        amt = abs(amount)
        company = str(row.get("company") or "")[:80]
        head = str(row.get("expense_head") or "Uncategorized")[:60]
        mode = str(row.get("mode") or "Bank")[:40]
        notes = str(row.get("notes") or "")[:200]
        direction = "in" if str(row.get("direction", "out")).lower() == "in" else "out"
        # duplicate check: same user + date + amount + similar company
        existing = await db.transactions.find_one({
            "user_id": user["id"], "date": d, "amount": amt,
            "company": company,
        }, {"_id": 0})
        doc = {
            "id": new_id(),
            "user_id": user["id"],
            "sr_no": next_sr,
            "order_index": 0,
            "date": d,
            "amount": amt,
            "mode": mode,
            "company": company,
            "vendor": company,
            "name": company,
            "expense_head": head,
            "head": head,
            "category": "income" if direction == "in" else "expense",
            "group": "",
            "direction": direction,
            "account": account,
            "notes": notes,
            "details": notes,
            "remarks": mode,
            "source": "upload",
            "created_at": now_iso(),
            "updated_at": now_iso(),
        }
        if existing and skip_duplicates:
            pending.append({**doc, "duplicate_of": existing["id"], "duplicate_existing": existing})
            continue
        await db.transactions.insert_one(dict(doc))
        inserted.append({k: v for k, v in doc.items() if k != "_id"})
        next_sr += 1
    return {"inserted": len(inserted), "transactions": inserted,
            "duplicate_count": len(pending), "duplicates": pending}


@api.get("/transactions/summary")
async def transactions_summary(month: Optional[str] = None, user=Depends(get_current_user)):
    month = month or datetime.now(timezone.utc).strftime("%Y-%m")
    q = {"user_id": user["id"], "date": {"$regex": f"^{month}"}}
    docs = await db.transactions.find(q, {"_id": 0}).to_list(10000)
    total_out = sum(d["amount"] for d in docs if d["direction"] == "out")
    total_in = sum(d["amount"] for d in docs if d["direction"] == "in")
    by_head: Dict[str, float] = {}
    for d in docs:
        if d["direction"] == "out":
            by_head[d["expense_head"]] = by_head.get(d["expense_head"], 0) + d["amount"]
    top5 = sorted(by_head.items(), key=lambda x: -x[1])[:5]
    # previous month
    try:
        y, m = [int(x) for x in month.split("-")]
        prev_m = m - 1 or 12
        prev_y = y if m != 1 else y - 1
        prev = f"{prev_y:04d}-{prev_m:02d}"
    except Exception:
        prev = month
    prev_docs = await db.transactions.find({"user_id": user["id"], "date": {"$regex": f"^{prev}"}}, {"_id": 0}).to_list(10000)
    prev_out = sum(d["amount"] for d in prev_docs if d["direction"] == "out")
    change = 0
    if prev_out:
        change = round(((total_out - prev_out) / prev_out) * 100)
    return {
        "month": month,
        "total_out": round(total_out, 2),
        "total_in": round(total_in, 2),
        "net": round(total_in - total_out, 2),
        "top_expense_heads": [{"head": h, "amount": round(a, 2)} for h, a in top5],
        "change_vs_prev_month_percent": change,
        "count": len(docs),
    }


# ───────────────────────────── investments ─────────────────────────────
@api.get("/investments", response_model=List[Investment])
async def list_investments(user=Depends(get_current_user)):
    docs = await db.investments.find({"user_id": user["id"]}, {"_id": 0}).sort("created_at", 1).to_list(1000)
    return [Investment(**d) for d in docs]


@api.post("/investments", response_model=Investment)
async def create_investment(body: InvestmentIn, user=Depends(get_current_user)):
    sr = await _next_sr("investments", user["id"])
    doc = {
        **body.model_dump(),
        "id": new_id(),
        "sr_no": sr,
        "user_id": user["id"],
        "created_at": now_iso(),
        "updated_at": now_iso(),
    }
    await db.investments.insert_one(dict(doc))
    return Investment(**doc)


@api.patch("/investments/{iid}", response_model=Investment)
async def update_investment(iid: str, body: Dict[str, Any], user=Depends(get_current_user)):
    allowed = {"kind", "type", "provider", "amount_invested", "start_date", "maturity_date",
               "rate_or_value", "current_value", "insured_for", "notes"}
    body = {k: v for k, v in body.items() if k in allowed}
    body["updated_at"] = now_iso()
    res = await db.investments.find_one_and_update(
        {"id": iid, "user_id": user["id"]}, {"$set": body}, projection={"_id": 0}, return_document=True
    )
    if not res:
        raise HTTPException(404, "Investment not found")
    return Investment(**res)


@api.delete("/investments/{iid}")
async def delete_investment(iid: str, user=Depends(get_current_user)):
    await db.investments.delete_one({"id": iid, "user_id": user["id"]})
    await _compact_sr("investments", user["id"])
    return {"ok": True}


@api.get("/investments/summary")
async def investments_summary(user=Depends(get_current_user)):
    docs = await db.investments.find({"user_id": user["id"]}, {"_id": 0}).to_list(1000)
    total_invested = sum(d["amount_invested"] for d in docs if (d.get("kind") or "investment") == "investment")
    total_insurance = sum(d["amount_invested"] for d in docs if (d.get("kind") or "investment") == "insurance")
    total_value = sum((d.get("current_value") or d["amount_invested"]) for d in docs)
    by_type: Dict[str, float] = {}
    for d in docs:
        by_type[d["type"]] = by_type.get(d["type"], 0) + (d.get("current_value") or d["amount_invested"])
    today = datetime.now(timezone.utc).date()
    upcoming = []
    for d in docs:
        md = d.get("maturity_date")
        if md:
            try:
                mdate = date.fromisoformat(md)
                diff = (mdate - today).days
                if 0 <= diff <= 90:
                    upcoming.append({"id": d["id"], "provider": d["provider"], "maturity_date": md, "days": diff})
            except Exception:
                pass
    return {
        "total_invested": round(total_invested, 2),
        "total_insurance": round(total_insurance, 2),
        "total_value": round(total_value, 2),
        "growth_percent": round(((total_value - total_invested) / total_invested) * 100, 2) if total_invested else 0,
        "allocation": [{"type": k, "value": round(v, 2)} for k, v in by_type.items()],
        "upcoming_maturities": upcoming,
        "count": len(docs),
        "investments_count": sum(1 for d in docs if (d.get("kind") or "investment") == "investment"),
        "insurance_count": sum(1 for d in docs if (d.get("kind") or "investment") == "insurance"),
    }


# ───────────────────────────── notes ─────────────────────────────
@api.get("/notes", response_model=List[Note])
async def list_notes(q: Optional[str] = None, tag: Optional[str] = None, project_id: Optional[str] = None, user=Depends(get_current_user)):
    query: Dict[str, Any] = {"user_id": user["id"]}
    if tag:
        query["tags"] = tag
    if q:
        query["$or"] = [
            {"title": {"$regex": q, "$options": "i"}},
            {"body": {"$regex": q, "$options": "i"}},
        ]
    if project_id:
        query["project_id"] = project_id
    docs = await db.notes.find(query, {"_id": 0}).sort("updated_at", -1).to_list(2000)
    return [Note(**d) for d in docs]


@api.post("/notes", response_model=Note)
async def create_note(body: NoteIn, user=Depends(get_current_user)):
    doc = {
        **body.model_dump(),
        "id": new_id(),
        "user_id": user["id"],
        "created_at": now_iso(),
        "updated_at": now_iso(),
    }
    if not doc.get("project_id"):
        doc["project_id"] = await _ensure_default_project(user["id"])
    await _assert_project_write(user["id"], doc["project_id"])
    await db.notes.insert_one(dict(doc))
    return Note(**doc)


@api.patch("/notes/{nid}", response_model=Note)
async def update_note(nid: str, body: Dict[str, Any], user=Depends(get_current_user)):
    allowed = {"title", "body", "tags", "pinned", "attachments", "flagged", "project_id"}
    body = {k: v for k, v in body.items() if k in allowed}
    body["updated_at"] = now_iso()
    res = await db.notes.find_one_and_update(
        {"id": nid, "user_id": user["id"]}, {"$set": body}, projection={"_id": 0}, return_document=True
    )
    if not res:
        raise HTTPException(404, "Note not found")
    return Note(**res)


@api.delete("/notes/{nid}")
async def delete_note(nid: str, user=Depends(get_current_user)):
    await db.notes.delete_one({"id": nid, "user_id": user["id"]})
    return {"ok": True}


class AppendListReq(BaseModel):
    # Either match by an existing note title (case-insensitive "contains") or by tag.
    title_hint: Optional[str] = None  # e.g. "shopping list"
    tag: Optional[str] = None         # e.g. "shopping"
    items: List[str]                  # new bullets to add
    create_if_missing: bool = True


@api.post("/notes/append-list", response_model=Note)
async def append_to_list(body: AppendListReq, user=Depends(get_current_user)):
    """Find an existing list-style note by title or tag; append items as bullets.
    If none match, create a new note."""
    q: Dict[str, Any] = {"user_id": user["id"]}
    existing = None
    if body.tag:
        tag = body.tag.lstrip("#").strip().lower()
        existing = await db.notes.find_one({**q, "tags": tag}, {"_id": 0})
    if not existing and body.title_hint:
        # case-insensitive contains on title
        existing = await db.notes.find_one(
            {**q, "title": {"$regex": body.title_hint, "$options": "i"}},
            {"_id": 0},
        )
    new_lines = [f"• {it.strip()}" for it in (body.items or []) if it and it.strip()]
    if not new_lines:
        raise HTTPException(400, "No items to append")

    if existing:
        existing_body = existing.get("body") or ""
        sep = "\n" if existing_body and not existing_body.endswith("\n") else ""
        new_body = existing_body + sep + "\n".join(new_lines)
        await db.notes.update_one(
            {"id": existing["id"]},
            {"$set": {"body": new_body, "updated_at": now_iso()}},
        )
        updated = await db.notes.find_one({"id": existing["id"]}, {"_id": 0})
        return Note(**updated)

    if not body.create_if_missing:
        raise HTTPException(404, "List not found")

    # Create new list note
    title = body.title_hint or (body.tag and f"#{body.tag.lstrip('#')}") or "New List"
    tags = []
    if body.tag:
        tags = [body.tag.lstrip("#").strip().lower()]
    doc = {
        "id": new_id(),
        "user_id": user["id"],
        "title": title,
        "body": "\n".join(new_lines),
        "tags": tags,
        "pinned": False,
        "created_at": now_iso(),
        "updated_at": now_iso(),
    }
    await db.notes.insert_one(dict(doc))
    return Note(**doc)


# ───────────────────────────── affirmations ─────────────────────────────
@api.get("/affirmations/today", response_model=Affirmation)
async def today_affirmation(user=Depends(get_current_user)):
    d = today_key()
    doc = await db.affirmations.find_one({"user_id": user["id"], "date": d}, {"_id": 0})
    if not doc:
        doc = {
            "id": new_id(),
            "user_id": user["id"],
            "date": d,
            "text": "",
            "updated_at": now_iso(),
        }
        await db.affirmations.insert_one(dict(doc))
    u = await db.users.find_one({"id": user["id"]}, {"_id": 0})
    doc["personal_fixed"] = (u or {}).get("personal_affirmation", "") or ""
    return Affirmation(**doc)


@api.put("/affirmations/today", response_model=Affirmation)
async def save_today_affirmation(body: AffirmationIn, user=Depends(get_current_user)):
    d = body.date or today_key()
    set_fields = {"updated_at": now_iso()}
    if body.text is not None:
        set_fields["text"] = body.text
    if body.personal_fixed is not None:
        # stored on user, not per-day (fixed)
        await db.users.update_one({"id": user["id"]},
                                  {"$set": {"personal_affirmation": body.personal_fixed}})
    await db.affirmations.update_one(
        {"user_id": user["id"], "date": d},
        {
            "$set": set_fields,
            "$setOnInsert": {"id": new_id(), "user_id": user["id"], "date": d},
        },
        upsert=True,
    )
    doc = await db.affirmations.find_one({"user_id": user["id"], "date": d}, {"_id": 0})
    u = await db.users.find_one({"id": user["id"]}, {"_id": 0})
    doc["personal_fixed"] = (u or {}).get("personal_affirmation", "") or ""
    return Affirmation(**doc)


# ───────────────────────────── dashboard ─────────────────────────────
@api.get("/dashboard/snapshot")
async def dashboard_snapshot(user=Depends(get_current_user)):
    today = today_key()
    tasks = await db.tasks.find({"user_id": user["id"], "status": "Pending"}, {"_id": 0}).to_list(1000)
    people: Dict[str, int] = {}
    for t in tasks:
        if t.get("name"):
            people[t["name"]] = people.get(t["name"], 0) + 1
    # cash today
    tx_today = await db.transactions.find({"user_id": user["id"], "date": today}, {"_id": 0}).to_list(1000)
    out_today = sum(t["amount"] for t in tx_today if t["direction"] == "out")
    in_today = sum(t["amount"] for t in tx_today if t["direction"] == "in")
    # loans
    loans = await db.loans.find({"user_id": user["id"]}, {"_id": 0}).to_list(5000)
    given = sum(l["amount"] for l in loans if l["status"] == "Given")
    taken = sum(l["amount"] for l in loans if l["status"] == "Taken")
    # routines
    summary = await routines_summary(user=user)
    # insights
    insights: List[str] = []
    # long-pending tasks
    five_days_ago = (datetime.now(timezone.utc) - timedelta(days=5)).date().isoformat()
    long_pending: Dict[str, int] = {}
    for t in tasks:
        if t.get("date", today) <= five_days_ago and t.get("name"):
            long_pending[t["name"]] = long_pending.get(t["name"], 0) + 1
    for name, c in long_pending.items():
        insights.append(f"{c} task{'s' if c != 1 else ''} from {name} pending for 5+ days")
    # overdue loans
    today_d = datetime.now(timezone.utc).date()
    for l in loans:
        if l["status"] in ("Given", "Pending") and l.get("repayment_date"):
            try:
                rd = date.fromisoformat(l["repayment_date"])
                over = (today_d - rd).days
                if over > 0:
                    insights.append(f"Loan to {l['name']} overdue by {over} days")
            except Exception:
                pass
    # routine shortfall
    if summary["total"] and summary["percent_today"] < 50:
        insights.append(f"You completed only {summary['percent_today']}% of routines today")
    # spending spike
    tx_sum = await transactions_summary(month=None, user=user)
    if tx_sum["change_vs_prev_month_percent"] and abs(tx_sum["change_vs_prev_month_percent"]) >= 15:
        sign = "up" if tx_sum["change_vs_prev_month_percent"] > 0 else "down"
        insights.append(f"Expenses this month are {sign} {abs(tx_sum['change_vs_prev_month_percent'])}% vs last month")

    return {
        "pending_tasks_count": len(tasks),
        "pending_by_person": [{"name": k, "count": v} for k, v in sorted(people.items(), key=lambda x: -x[1])],
        "routine_percent_today": summary["percent_today"],
        "cash_out_today": round(out_today, 2),
        "cash_in_today": round(in_today, 2),
        "net_cash_today": round(in_today - out_today, 2),
        "loans_net_exposure": round(given - taken, 2),
        "insights": insights[:6],
    }


# ───────────────────────────── AI ─────────────────────────────
class ChatReq(BaseModel):
    message: str
    session_id: Optional[str] = None


@api.post("/ai/chat")
async def ai_chat(body: ChatReq, user=Depends(get_current_user)):
    if not EMERGENT_LLM_KEY:
        raise HTTPException(500, "LLM not configured")
    session_id = body.session_id or f"mm-{user['id']}"
    # gather lightweight context
    tasks = await db.tasks.find({"user_id": user["id"]}, {"_id": 0}).to_list(200)
    loans = await db.loans.find({"user_id": user["id"]}, {"_id": 0}).to_list(200)
    tx = await db.transactions.find({"user_id": user["id"]}, {"_id": 0}).sort("date", -1).to_list(200)
    routines = await db.routines.find({"user_id": user["id"]}, {"_id": 0}).to_list(200)
    import json
    context = {
        "tasks": tasks[:50],
        "loans": loans[:50],
        "recent_transactions": tx[:50],
        "routines": routines[:50],
    }
    system = (
        "You are Mind Matters — a calm, intelligent personal operating system assistant. "
        "Reply concisely (1-4 short sentences unless the user asks for detail). "
        "You have access to the user's data as JSON. Use it to answer questions about "
        "tasks, loans, cash flow, routines, investments, and notes. Use INR (₹) where relevant."
    )
    chat = LlmChat(api_key=EMERGENT_LLM_KEY, session_id=session_id, system_message=system)\
        .with_model("gemini", "gemini-3-flash-preview")
    msg = UserMessage(text=f"User data context:\n{json.dumps(context, default=str)[:12000]}\n\nUser question: {body.message}")
    reply = await chat.send_message(msg)
    # persist
    await db.chat_messages.insert_one({
        "id": new_id(),
        "user_id": user["id"],
        "session_id": session_id,
        "role": "user",
        "text": body.message,
        "created_at": now_iso(),
    })
    await db.chat_messages.insert_one({
        "id": new_id(),
        "user_id": user["id"],
        "session_id": session_id,
        "role": "assistant",
        "text": reply,
        "created_at": now_iso(),
    })
    return {"reply": reply, "session_id": session_id}


class ParseReq(BaseModel):
    text: str
    kind: Literal["task", "expense", "note", "auto"] = "auto"


@api.post("/ai/parse")
async def ai_parse(body: ParseReq, user=Depends(get_current_user)):
    """Natural language → structured object for task/expense/note."""
    if not EMERGENT_LLM_KEY:
        raise HTTPException(500, "LLM not configured")
    import json
    system = (
        "Extract structured data from a short natural-language instruction. "
        "Return ONLY JSON. Schemas:\n"
        "task: {kind:'task', name:string (responsible person or ''), task:string, details:string, date:'YYYY-MM-DD' or null, status:'Pending'}\n"
        "expense: {kind:'expense', amount:number, expense_head:string, company:string, direction:'out'|'in', date:'YYYY-MM-DD' or null, mode:string, notes:string}\n"
        "note: {kind:'note', title:string, body:string, tags:[string]}\n"
    )
    prompt = f"Kind hint: {body.kind}. Instruction: {body.text}"
    chat = LlmChat(api_key=EMERGENT_LLM_KEY, session_id=f"parse-{new_id()}", system_message=system)\
        .with_model("gemini", "gemini-3-flash-preview")
    resp = await chat.send_message(UserMessage(text=prompt))
    text = resp.strip()
    if text.startswith("```"):
        text = text.strip("`")
        if text.startswith("json"):
            text = text[4:]
    try:
        data = json.loads(text)
        return {"parsed": data, "raw": resp}
    except Exception:
        return {"parsed": None, "raw": resp}


# ───────────────────────────── daily quote (replaces news) ─────────────────────────────
@api.get("/quote/today")
async def quote_today():
    """Fetch a single inspirational quote from a public, free API."""
    # Use ZenQuotes — no key needed. Fallback to a static list on failure.
    try:
        async with httpx.AsyncClient(timeout=6) as cli:
            r = await cli.get("https://zenquotes.io/api/today")
            arr = r.json()
            if isinstance(arr, list) and arr:
                q = arr[0]
                text = q.get("q") or ""
                author = q.get("a") or "Unknown"
                if text:
                    return {"text": text.strip(), "author": author.strip()}
    except Exception as e:
        logger.warning(f"quote fetch failed: {e}")
    fallback = [
        ("Discipline is the bridge between goals and accomplishment.", "Jim Rohn"),
        ("The secret of getting ahead is getting started.", "Mark Twain"),
        ("Small daily improvements over time lead to stunning results.", "Robin Sharma"),
        ("You don't have to be great to start, but you have to start to be great.", "Zig Ziglar"),
    ]
    # pick by day-of-year so it's deterministic per day
    idx = datetime.now(timezone.utc).timetuple().tm_yday % len(fallback)
    text, author = fallback[idx]
    return {"text": text, "author": author}


# ───────────────────────────── news (GNews) ─────────────────────────────
NEWS_API_KEY = os.environ.get("NEWS_API_KEY", "")


@api.get("/news/headlines")
async def headlines():
    if NEWS_API_KEY:
        try:
            async with httpx.AsyncClient(timeout=8) as cli:
                r = await cli.get(
                    "https://gnews.io/api/v4/top-headlines",
                    params={"lang": "en", "country": "in", "max": 5,
                            "apikey": NEWS_API_KEY, "category": "general"},
                )
                data = r.json()
            arts = data.get("articles", []) or []
            return {"headlines": [
                {"title": a.get("title", ""), "source": (a.get("source") or {}).get("name", ""),
                 "url": a.get("url", "")} for a in arts[:3]
            ]}
        except Exception as e:
            logger.warning(f"GNews failed: {e}")
    # fallback
    return {
        "headlines": [
            {"title": "Global markets steady as central banks hold rates", "source": "Reuters"},
            {"title": "Tech giants report strong AI-driven earnings growth", "source": "Bloomberg"},
            {"title": "India's economy projected to grow 6.8% next fiscal year", "source": "Financial Times"},
        ]
    }


# ───────────────────────────── weather ─────────────────────────────
@api.get("/weather")
async def weather(lat: Optional[float] = None, lon: Optional[float] = None):
    """Open-Meteo free API, no key required (async httpx)."""
    try:
        if lat is None or lon is None:
            lat, lon = 19.0760, 72.8777  # Mumbai default
        async with httpx.AsyncClient(timeout=6) as cli:
            r = await cli.get(
                "https://api.open-meteo.com/v1/forecast",
                params={"latitude": lat, "longitude": lon, "current": "temperature_2m,weather_code"},
            )
            data = r.json()
        cur = data.get("current", {})
        code = cur.get("weather_code", 0)
        CODES = {
            0: "Clear", 1: "Mainly clear", 2: "Partly cloudy", 3: "Overcast",
            45: "Fog", 48: "Fog", 51: "Light drizzle", 53: "Drizzle", 55: "Drizzle",
            61: "Light rain", 63: "Rain", 65: "Heavy rain",
            71: "Snow", 73: "Snow", 75: "Heavy snow",
            80: "Rain showers", 81: "Rain showers", 82: "Rain showers",
            95: "Thunderstorm", 96: "Thunderstorm", 99: "Thunderstorm",
        }
        return {
            "temperature": cur.get("temperature_2m"),
            "label": CODES.get(code, "—"),
            "code": code,
            "lat": lat,
            "lon": lon,
        }
    except Exception as e:
        logger.warning(f"weather failed: {e}")
        return {"temperature": None, "label": "—", "code": 0}


# ═════════════════════════════════════════════════════════════════════
# DOCUMENTS / REMINDERS / TELEGRAM / SHARE / SETTINGS  (Phase 2+)
# ═════════════════════════════════════════════════════════════════════

# --- helper: expand "bill_to.name" style keys into a nested dict ---
def _unflatten(d: Dict[str, Any]) -> Dict[str, Any]:
    out: Dict[str, Any] = {}
    for k, v in d.items():
        parts = k.split(".")
        cur = out
        for p in parts[:-1]:
            cur = cur.setdefault(p, {})
        cur[parts[-1]] = v
    return out


# --- internal creators reused from Telegram + share endpoints ---
async def _create_task_internal(user_id: str, parsed: Dict[str, Any]):
    sr = await _next_sr("tasks", user_id)
    doc = {
        "id": new_id(), "sr_no": sr, "user_id": user_id,
        "date": parsed.get("date") or today_key(),
        "name": parsed.get("name") or "",
        "task": parsed.get("task") or "",
        "details": parsed.get("details") or "",
        "status": parsed.get("status") or "Pending",
        "created_at": now_iso(), "updated_at": now_iso(),
    }
    await db.tasks.insert_one(dict(doc))
    return doc


async def _create_expense_internal(user_id: str, parsed: Dict[str, Any]):
    doc = {
        "id": new_id(), "user_id": user_id,
        "date": (parsed.get("date") or today_key()),
        "amount": float(parsed.get("amount") or 0),
        "mode": parsed.get("mode") or "Cash",
        "company": parsed.get("company") or "",
        "expense_head": parsed.get("expense_head") or "Uncategorized",
        "direction": parsed.get("direction") or "out",
        "account": "Personal",
        "notes": parsed.get("notes") or "",
        "source": "telegram",
        "created_at": now_iso(), "updated_at": now_iso(),
    }
    await db.transactions.insert_one(dict(doc))
    return doc


async def _create_note_internal(user_id: str, parsed: Dict[str, Any]):
    doc = {
        "id": new_id(), "user_id": user_id,
        "title": parsed.get("title") or "",
        "body": parsed.get("body") or "",
        "tags": parsed.get("tags") or [],
        "pinned": False,
        "created_at": now_iso(), "updated_at": now_iso(),
    }
    await db.notes.insert_one(dict(doc))
    return doc


async def _ai_parse_text(text: str, kind: str = "auto") -> Dict[str, Any]:
    if not EMERGENT_LLM_KEY:
        return {}
    today = today_key()
    system = (
        "Extract structured data from a short natural-language instruction. Return ONLY JSON.\n"
        f"TODAY'S DATE is {today} — use it for relative phrases (today, tomorrow, last week).\n"
        "task: {kind:'task', name, task, details, date:'YYYY-MM-DD' or null, status:'Pending'}\n"
        "expense: {kind:'expense', amount:number, expense_head, company, direction:'out'|'in', date, mode, notes}\n"
        "note: {kind:'note', title, body, tags:[string]}\n"
    )
    chat = LlmChat(api_key=EMERGENT_LLM_KEY, session_id=f"p-{new_id()}",
                   system_message=system).with_model("gemini", "gemini-3-flash-preview")
    resp = await chat.send_message(UserMessage(text=f"Kind: {kind}. Instruction: {text}"))
    t = resp.strip()
    if t.startswith("```"):
        t = t.strip("`")
        if t.startswith("json"):
            t = t[4:]
    try:
        return json.loads(t)
    except Exception:
        return {}


# ───────────────────── Documents / Templates ─────────────────────
class GenerateDocReq(BaseModel):
    template_id: str
    data: Dict[str, Any]


@api.get("/documents/templates")
async def list_templates(user=Depends(get_current_user)):
    # built-in + user custom
    user_templates = await db.templates.find(
        {"user_id": user["id"]}, {"_id": 0}
    ).sort("created_at", -1).to_list(200)
    return {"templates": BUILTIN_TEMPLATES + user_templates}


@api.post("/documents/templates/upload")
async def upload_template(
    file: UploadFile = File(...),
    name: str = Form(...),
    kind: str = Form("invoice"),
    user=Depends(get_current_user),
):
    """Upload a .docx with {{placeholders}}. Returns the template meta."""
    if not (file.filename or "").lower().endswith(".docx"):
        raise HTTPException(400, "Only .docx templates are supported right now")
    content = await file.read()
    # Quick placeholder discovery
    try:
        from docx import Document
        d = Document(io.BytesIO(content))
        text_all = "\n".join(p.text for p in d.paragraphs)
        for tbl in d.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    text_all += "\n" + cell.text
    except Exception as e:
        raise HTTPException(400, f"Could not parse docx: {e}")
    import re
    placeholders = sorted(set(re.findall(r"{{\s*([a-zA-Z0-9_\.]+)\s*}}", text_all)))
    tmpl = {
        "id": new_id(),
        "user_id": user["id"],
        "name": name,
        "kind": kind,
        "builtin": False,
        "file_name": file.filename,
        "content_b64": __import__("base64").b64encode(content).decode("ascii"),
        "required_fields": [{"key": p, "label": p.replace("_", " ").title(), "type": "text"}
                            for p in placeholders],
        "created_at": now_iso(),
    }
    await db.templates.insert_one(dict(tmpl))
    # strip content from response
    resp = {k: v for k, v in tmpl.items() if k != "content_b64"}
    return resp


@api.delete("/documents/templates/{tid}")
async def delete_template(tid: str, user=Depends(get_current_user)):
    r = await db.templates.delete_one({"id": tid, "user_id": user["id"]})
    if r.deleted_count == 0:
        raise HTTPException(404, "Template not found")
    return {"ok": True}


def _render_template_bytes(template_id: str, data: Dict[str, Any], db_template: Optional[dict]) -> tuple[str, bytes]:
    """Returns (filename, bytes)."""
    data = _unflatten(data or {})
    # built-in
    if template_id in ("rkm_donation_receipt", "krm_huf_invoice"):
        content = render_by_template_id(template_id, data)
        name = "receipt.pdf" if template_id.endswith("receipt") else "invoice.pdf"
        return name, content
    # user-uploaded docx
    if db_template and db_template.get("content_b64"):
        import base64
        from docxtpl import DocxTemplate
        raw = base64.b64decode(db_template["content_b64"])
        doc = DocxTemplate(io.BytesIO(raw))
        doc.render(data)
        out = io.BytesIO()
        doc.save(out)
        return f"{db_template.get('name','document')}.docx", out.getvalue()
    raise HTTPException(404, "Template not found")


@api.post("/documents/generate")
async def generate_document(body: GenerateDocReq, user=Depends(get_current_user)):
    tpl = None
    if body.template_id not in ("rkm_donation_receipt", "krm_huf_invoice"):
        tpl = await db.templates.find_one(
            {"id": body.template_id, "user_id": user["id"]}, {"_id": 0}
        )
        if not tpl:
            raise HTTPException(404, "Template not found")
    filename, content = _render_template_bytes(body.template_id, body.data, tpl)
    from fastapi.responses import Response
    mime = "application/pdf" if filename.endswith(".pdf") else \
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    return Response(
        content=content,
        media_type=mime,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


class ShareDocReq(BaseModel):
    template_id: str
    data: Dict[str, Any]
    caption: Optional[str] = ""


@api.post("/documents/share-telegram")
async def share_document_telegram(body: ShareDocReq, user=Depends(get_current_user)):
    cid = user.get("telegram_chat_id")
    if not cid:
        raise HTTPException(400, "Telegram not linked. Connect in Settings.")
    tpl = None
    if body.template_id not in ("rkm_donation_receipt", "krm_huf_invoice"):
        tpl = await db.templates.find_one(
            {"id": body.template_id, "user_id": user["id"]}, {"_id": 0}
        )
        if not tpl:
            raise HTTPException(404, "Template not found")
    filename, content = _render_template_bytes(body.template_id, body.data, tpl)
    res = await tg_send_document(cid, filename, content, caption=body.caption or "")
    return {"ok": bool(res and res.get("ok")), "telegram": res}


# ───────────────────── Universal search ─────────────────────
@api.get("/search")
async def universal_search(q: str = "", user=Depends(get_current_user)):
    """Search across tasks, notes, transactions, reminders, decisions, vault."""
    q = (q or "").strip()
    if not q or len(q) < 1:
        return {"groups": []}
    import re as _re
    regex = {"$regex": _re.escape(q), "$options": "i"}
    groups = []

    # Tasks
    docs = await db.tasks.find(
        {"user_id": user["id"],
         "$or": [{"task": regex}, {"details": regex}, {"name": regex}, {"group": regex}]},
        {"_id": 0, "id": 1, "task": 1, "details": 1, "name": 1, "status": 1, "date": 1, "group": 1},
    ).limit(8).to_list(8)
    if docs:
        groups.append({
            "module": "tasks", "label": "Tasks", "route": "/tasks",
            "items": [{
                "id": d["id"],
                "title": d.get("task") or d.get("details") or "Task",
                "snippet": " · ".join(filter(None, [d.get("name"), d.get("details"), d.get("status"), d.get("date")]))[:160],
            } for d in docs],
        })

    # Notes
    docs = await db.notes.find(
        {"user_id": user["id"],
         "$or": [{"title": regex}, {"body": regex}]},
        {"_id": 0, "id": 1, "title": 1, "body": 1, "tags": 1},
    ).limit(8).to_list(8)
    if docs:
        groups.append({
            "module": "notes", "label": "Notes", "route": "/notes",
            "items": [{
                "id": d["id"],
                "title": d.get("title") or "(untitled)",
                "snippet": (d.get("body") or "")[:160],
            } for d in docs],
        })

    # Transactions
    docs = await db.transactions.find(
        {"user_id": user["id"],
         "$or": [{"vendor": regex}, {"name": regex}, {"details": regex}, {"notes": regex}, {"head": regex}]},
        {"_id": 0, "id": 1, "vendor": 1, "name": 1, "details": 1, "amount": 1, "head": 1, "date": 1},
    ).limit(8).to_list(8)
    if docs:
        groups.append({
            "module": "transactions", "label": "Cash Flow", "route": "/cash-flow",
            "items": [{
                "id": d["id"],
                "title": d.get("vendor") or d.get("name") or "Entry",
                "snippet": " · ".join(filter(None, [
                    f"₹{d.get('amount')}" if d.get("amount") else "",
                    d.get("head"), d.get("details"), d.get("date"),
                ]))[:160],
            } for d in docs],
        })

    # Reminders
    docs = await db.reminders.find(
        {"user_id": user["id"],
         "$or": [{"title": regex}, {"notes": regex}]},
        {"_id": 0, "id": 1, "title": 1, "notes": 1, "fire_at": 1},
    ).limit(8).to_list(8)
    if docs:
        groups.append({
            "module": "reminders", "label": "Reminders", "route": "/reminders",
            "items": [{
                "id": d["id"],
                "title": d.get("title") or "Reminder",
                "snippet": " · ".join(filter(None, [d.get("notes"), d.get("fire_at", "")[:16]]))[:160],
            } for d in docs],
        })

    return {"groups": groups}


# ───────────────────── Telegram setup PDF ─────────────────────
@api.get("/docs/telegram-setup.pdf")
async def telegram_setup_pdf():
    """Public — step-by-step PDF guide for linking a private Telegram bot."""
    headers = ["Step", "What to do"]
    rows = [
        ["1", "Open Telegram and search for @BotFather. Open the chat and tap Start."],
        ["2", "Send /newbot. BotFather will ask for a display name (e.g. \"Mind Matters\")."],
        ["3", "Choose a unique username ending in 'bot' (e.g. mindmatters_karan_bot)."],
        ["4", "BotFather replies with an HTTP API token like 7123456789:AAH...XYZ. Copy it."],
        ["5", "In Mind Matters → Settings → Connect Telegram, paste the token (or send to admin) and Save."],
        ["6", "Tap 'Generate code' in Settings — a deep-link button appears."],
        ["7", "Tap 'Open in Telegram'. Inside the chat, tap Start. Your account is now linked."],
        ["8", "From now on, send any text or receipt photo to your bot. AI will parse and confirm."],
        ["9", "Reminders also auto-ping you on Telegram at the scheduled time."],
        ["10", "To unlink later, return to Settings → Connect Telegram → Unlink."],
    ]
    pdf = render_simple_statement(
        title="Telegram Bot — Setup Guide",
        subtitle="Mind Matters · 2-way private bot in 10 steps",
        meta={"App": "Mind Matters", "Version": "v2.2", "Last updated": today_key()},
        table_headers=headers,
        table_rows=rows,
        summary_rows=[
            ("Tip", "Each user runs their own private bot — your data never crosses accounts."),
            ("Need help?", "Settings → Connect Telegram has live status + test-ping button."),
        ],
        footer="© Mind Matters — Personal Operating System",
    )
    from fastapi.responses import Response
    return Response(
        content=pdf,
        media_type="application/pdf",
        headers={"Content-Disposition": 'attachment; filename="mind-matters-telegram-setup.pdf"'},
    )


# ───────────────────── Projects + Sharing + Comments (v2.17) ─────────────────────
# Every data row (tasks, routines, transactions, notes, reminders, deadlines)
# carries a `project_id`. Users can own multiple projects and be invited to
# others. Roles: admin | editor | commenter | viewer.

_PROJECT_DATA_COLLECTIONS = (
    "tasks", "routines", "transactions", "notes", "reminders", "deadlines",
)

_PROJECT_COLORS = [
    "#C9A961", "#7AB8FF", "#9EE493", "#FF9F7A", "#D89BFF",
    "#FFD27A", "#7AF0D8", "#FF8AB3", "#A7F39B", "#B5A8FF",
]


class ProjectIn(BaseModel):
    name: str
    color: Optional[str] = None


class Project(BaseModel):
    model_config = ConfigDict(extra="allow")
    id: str
    owner_id: str
    name: str
    color: str = "#C9A961"
    is_default: bool = False
    created_at: str
    # Hydrated server-side (not persisted):
    role: Optional[str] = None  # admin | editor | commenter | viewer
    shared: bool = False
    member_count: int = 1


class ProjectMember(BaseModel):
    id: str
    project_id: str
    user_id: Optional[str] = None      # populated once invitee creates / has an account
    invited_email: str
    role: Literal["admin", "editor", "commenter", "viewer"]
    accepted: bool = False
    invited_by: str
    created_at: str


class ShareReq(BaseModel):
    email: str
    role: Literal["admin", "editor", "commenter", "viewer"] = "editor"


class MemberRoleReq(BaseModel):
    role: Literal["admin", "editor", "commenter", "viewer"]


class CommentIn(BaseModel):
    resource_type: Literal["task", "routine", "transaction", "note"]
    resource_id: str
    body: str


class Comment(BaseModel):
    id: str
    project_id: str
    resource_type: str
    resource_id: str
    user_id: str
    user_name: str = ""
    body: str
    created_at: str


async def _ensure_default_project(user_id: str) -> str:
    """Return the user's default project id; create one if missing.

    First-run: if user has any existing data rows but no project, create the
    default project AND back-fill `project_id` on every existing row.
    """
    existing = await db.projects.find_one(
        {"owner_id": user_id, "is_default": True}, {"_id": 0}
    )
    if existing:
        return existing["id"]

    # Maybe user has a non-default project — use the oldest one as default.
    any_proj = await db.projects.find_one(
        {"owner_id": user_id}, {"_id": 0}, sort=[("created_at", 1)]
    )
    if any_proj:
        await db.projects.update_one(
            {"id": any_proj["id"]}, {"$set": {"is_default": True}}
        )
        return any_proj["id"]

    proj_id = new_id()
    doc = {
        "id": proj_id, "owner_id": user_id,
        "name": "Personal", "color": "#C9A961",
        "is_default": True, "created_at": now_iso(),
    }
    await db.projects.insert_one(dict(doc))
    # Back-fill project_id on any existing rows for this user.
    for coll in _PROJECT_DATA_COLLECTIONS:
        await db[coll].update_many(
            {"user_id": user_id, "$or": [{"project_id": {"$exists": False}}, {"project_id": None}]},
            {"$set": {"project_id": proj_id}},
        )
    return proj_id


async def _user_accessible_project_ids(user_id: str, email: str = "") -> List[str]:
    owned = await db.projects.find({"owner_id": user_id}, {"_id": 0, "id": 1}).to_list(500)
    pids = [p["id"] for p in owned]
    member_query: Dict[str, Any] = {"user_id": user_id, "accepted": True}
    members = await db.project_members.find(member_query, {"_id": 0, "project_id": 1}).to_list(500)
    pids.extend(m["project_id"] for m in members)
    # Also accept invites pending acceptance by email (auto-accept on next access)
    if email:
        pending = await db.project_members.find(
            {"invited_email": email.lower(), "accepted": False},
            {"_id": 0, "id": 1, "project_id": 1},
        ).to_list(200)
        if pending:
            ids = [p["id"] for p in pending]
            await db.project_members.update_many(
                {"id": {"$in": ids}},
                {"$set": {"user_id": user_id, "accepted": True, "accepted_at": now_iso()}},
            )
            pids.extend(p["project_id"] for p in pending)
    return list(dict.fromkeys(pids))  # preserve order, dedupe


async def _user_role_in_project(user_id: str, project_id: str) -> Optional[str]:
    proj = await db.projects.find_one({"id": project_id}, {"_id": 0})
    if not proj:
        return None
    if proj["owner_id"] == user_id:
        return "admin"
    member = await db.project_members.find_one(
        {"project_id": project_id, "user_id": user_id, "accepted": True}, {"_id": 0}
    )
    return member["role"] if member else None


async def _assert_project_write(user_id: str, project_id: Optional[str]):
    """Raise 403 if user lacks editor/admin access to the given project."""
    if not project_id:
        return
    role = await _user_role_in_project(user_id, project_id)
    if role not in ("admin", "editor"):
        raise HTTPException(403, "You don't have write access to this project")


async def _hydrate_project_for_user(proj: Dict[str, Any], user_id: str) -> Dict[str, Any]:
    role = await _user_role_in_project(user_id, proj["id"])
    member_count = 1 + await db.project_members.count_documents(
        {"project_id": proj["id"], "accepted": True}
    )
    return {
        **{k: v for k, v in proj.items() if k != "_id"},
        "role": role or "viewer",
        "shared": proj.get("owner_id") != user_id or member_count > 1,
        "member_count": member_count,
    }


@api.get("/projects", response_model=List[Project])
async def list_projects(user=Depends(get_current_user)):
    await _ensure_default_project(user["id"])
    pids = await _user_accessible_project_ids(user["id"], user.get("email", ""))
    docs = await db.projects.find({"id": {"$in": pids}}, {"_id": 0}).to_list(500)
    out = []
    for d in docs:
        out.append(await _hydrate_project_for_user(d, user["id"]))
    out.sort(key=lambda p: (not p.get("is_default"), p.get("name", "").lower()))
    return [Project(**p) for p in out]


@api.post("/projects", response_model=Project)
async def create_project(body: ProjectIn, user=Depends(get_current_user)):
    name = body.name.strip() or "Project"
    color = (body.color or "").strip() or _PROJECT_COLORS[
        len(await db.projects.find({"owner_id": user["id"]}, {"_id": 0, "id": 1}).to_list(50)) % len(_PROJECT_COLORS)
    ]
    doc = {
        "id": new_id(), "owner_id": user["id"], "name": name, "color": color,
        "is_default": False, "created_at": now_iso(),
    }
    await db.projects.insert_one(dict(doc))
    return Project(**await _hydrate_project_for_user(doc, user["id"]))


@api.patch("/projects/{pid}", response_model=Project)
async def update_project(pid: str, body: Dict[str, Any], user=Depends(get_current_user)):
    proj = await db.projects.find_one({"id": pid}, {"_id": 0})
    if not proj:
        raise HTTPException(404, "Project not found")
    role = await _user_role_in_project(user["id"], pid)
    if role != "admin":
        raise HTTPException(403, "Only admins can edit a project")
    allowed = {"name", "color"}
    body = {k: v for k, v in body.items() if k in allowed}
    if body:
        await db.projects.update_one({"id": pid}, {"$set": body})
    fresh = await db.projects.find_one({"id": pid}, {"_id": 0})
    return Project(**await _hydrate_project_for_user(fresh, user["id"]))


@api.delete("/projects/{pid}")
async def delete_project(pid: str, user=Depends(get_current_user)):
    proj = await db.projects.find_one({"id": pid}, {"_id": 0})
    if not proj:
        raise HTTPException(404, "Project not found")
    if proj["owner_id"] != user["id"]:
        raise HTTPException(403, "Only the owner can delete a project")
    if proj.get("is_default"):
        raise HTTPException(400, "Cannot delete the default project")
    # Cascade: remove members, comments, and reassign rows to default project.
    default_id = await _ensure_default_project(user["id"])
    for coll in _PROJECT_DATA_COLLECTIONS:
        await db[coll].update_many(
            {"user_id": user["id"], "project_id": pid},
            {"$set": {"project_id": default_id}},
        )
    await db.project_members.delete_many({"project_id": pid})
    await db.comments.delete_many({"project_id": pid})
    await db.projects.delete_one({"id": pid})
    return {"ok": True}


@api.get("/projects/{pid}/members")
async def list_project_members(pid: str, user=Depends(get_current_user)):
    role = await _user_role_in_project(user["id"], pid)
    if not role:
        raise HTTPException(403, "Not a member of this project")
    proj = await db.projects.find_one({"id": pid}, {"_id": 0})
    owner = await db.users.find_one({"id": proj["owner_id"]}, {"_id": 0, "id": 1, "email": 1, "first_name": 1})
    members = await db.project_members.find({"project_id": pid}, {"_id": 0}).to_list(200)
    return {
        "owner": {
            "user_id": owner["id"] if owner else proj["owner_id"],
            "email": owner.get("email") if owner else "",
            "first_name": owner.get("first_name", "") if owner else "",
            "role": "admin",
        },
        "members": members,
    }


@api.post("/projects/{pid}/share")
async def share_project(pid: str, body: ShareReq, user=Depends(get_current_user)):
    role = await _user_role_in_project(user["id"], pid)
    if role != "admin":
        raise HTTPException(403, "Only admins can invite members")
    email = body.email.strip().lower()
    if not email or "@" not in email:
        raise HTTPException(400, "Invalid email")
    proj = await db.projects.find_one({"id": pid}, {"_id": 0})
    if not proj:
        raise HTTPException(404, "Project not found")
    if proj["owner_id"] == user["id"] and user.get("email") == email:
        raise HTTPException(400, "Cannot invite yourself")
    invitee = await db.users.find_one({"email": email}, {"_id": 0})
    existing = await db.project_members.find_one(
        {"project_id": pid, "invited_email": email}, {"_id": 0}
    )
    if existing:
        await db.project_members.update_one(
            {"id": existing["id"]}, {"$set": {"role": body.role}}
        )
        return {"ok": True, "updated": True}
    member = {
        "id": new_id(), "project_id": pid,
        "user_id": invitee["id"] if invitee else None,
        "invited_email": email, "role": body.role,
        "accepted": bool(invitee), "invited_by": user["id"],
        "created_at": now_iso(),
    }
    if invitee:
        member["accepted_at"] = now_iso()
    await db.project_members.insert_one(dict(member))
    return {"ok": True, "member_id": member["id"], "accepted": member["accepted"]}


@api.patch("/projects/{pid}/members/{mid}")
async def update_project_member(pid: str, mid: str, body: MemberRoleReq,
                                user=Depends(get_current_user)):
    role = await _user_role_in_project(user["id"], pid)
    if role != "admin":
        raise HTTPException(403, "Only admins can change roles")
    res = await db.project_members.update_one(
        {"id": mid, "project_id": pid}, {"$set": {"role": body.role}}
    )
    if res.matched_count == 0:
        raise HTTPException(404, "Member not found")
    return {"ok": True}


@api.delete("/projects/{pid}/members/{mid}")
async def remove_project_member(pid: str, mid: str, user=Depends(get_current_user)):
    role = await _user_role_in_project(user["id"], pid)
    if role != "admin":
        # Allow a member to remove themselves (leave a project)
        member = await db.project_members.find_one(
            {"id": mid, "project_id": pid}, {"_id": 0}
        )
        if not member or member.get("user_id") != user["id"]:
            raise HTTPException(403, "Only admins can remove members")
    await db.project_members.delete_one({"id": mid, "project_id": pid})
    return {"ok": True}


@api.get("/projects/{pid}/comments", response_model=List[Comment])
async def list_comments(pid: str, resource_type: Optional[str] = None,
                        resource_id: Optional[str] = None,
                        user=Depends(get_current_user)):
    role = await _user_role_in_project(user["id"], pid)
    if not role:
        raise HTTPException(403, "Not a member of this project")
    q: Dict[str, Any] = {"project_id": pid}
    if resource_type:
        q["resource_type"] = resource_type
    if resource_id:
        q["resource_id"] = resource_id
    docs = await db.comments.find(q, {"_id": 0}).sort("created_at", 1).to_list(500)
    return [Comment(**d) for d in docs]


@api.post("/projects/{pid}/comments", response_model=Comment)
async def post_comment(pid: str, body: CommentIn, user=Depends(get_current_user)):
    role = await _user_role_in_project(user["id"], pid)
    if role not in ("admin", "editor", "commenter"):
        raise HTTPException(403, "You can't comment in this project")
    doc = {
        "id": new_id(), "project_id": pid,
        "resource_type": body.resource_type, "resource_id": body.resource_id,
        "user_id": user["id"],
        "user_name": user.get("first_name", "") or user.get("email", "")[:24],
        "body": body.body.strip(),
        "created_at": now_iso(),
    }
    if not doc["body"]:
        raise HTTPException(400, "Comment body required")
    await db.comments.insert_one(dict(doc))
    return Comment(**doc)


@api.delete("/comments/{cid}")
async def delete_comment(cid: str, user=Depends(get_current_user)):
    c = await db.comments.find_one({"id": cid}, {"_id": 0})
    if not c:
        raise HTTPException(404, "Comment not found")
    role = await _user_role_in_project(user["id"], c["project_id"])
    if c["user_id"] != user["id"] and role != "admin":
        raise HTTPException(403, "Cannot delete this comment")
    await db.comments.delete_one({"id": cid})
    return {"ok": True}


# ───────────────────── v2.17 startup migrations ─────────────────────
_LEGACY_DEMO_SIGNATURES = {
    "tasks": [
        {"task": "call about repair", "name": "Brinda"},
    ],
    "routines": [
        {"activity": "20 min walk", "name": "Self"},
    ],
    "transactions": [
        {"vendor": "Coffee Shop", "details": "morning latte"},
    ],
    "notes": [
        {"title": "Shopping List", "body": "• milk\n• eggs\n• bread"},
    ],
    "reminders": [
        {"title": "Welcome to Mind Matters"},
    ],
    "deadlines": [
        {"title": "Quarterly Review"},
    ],
}


async def _purge_legacy_demo_data() -> None:
    """Remove all rows that match the OLD demo-seed signatures (Item 47).

    Also fully wipes the unattached `you@mindmatters.local` demo user's data
    if it has never been claimed by a real signup.
    """
    total = 0
    for coll, sigs in _LEGACY_DEMO_SIGNATURES.items():
        for sig in sigs:
            res = await db[coll].delete_many(sig)
            total += res.deleted_count
    # Demo user without password — wipe its data entirely.
    demo = await db.users.find_one(
        {"email": "you@mindmatters.local", "password_hash": {"$in": [None, ""]}},
        {"_id": 0, "id": 1},
    )
    if demo:
        for coll in _PROJECT_DATA_COLLECTIONS + ("routine_logs", "affirmations"):
            res = await db[coll].delete_many({"user_id": demo["id"]})
            total += res.deleted_count
    if total:
        logger.info(f"v2.17 demo purge: removed {total} legacy demo rows")


async def _backfill_default_projects() -> None:
    """Ensure every user has a default 'Personal' project + back-fill project_id."""
    users = await db.users.find({}, {"_id": 0, "id": 1}).to_list(5000)
    fixed = 0
    for u in users:
        before = await db.projects.count_documents({"owner_id": u["id"]})
        await _ensure_default_project(u["id"])
        if not before:
            fixed += 1
    if fixed:
        logger.info(f"v2.17 project backfill: seeded default project for {fixed} users")



class ResetReq(BaseModel):
    confirm: str  # must equal "RESET"


@api.post("/reset/seed")
async def reset_and_seed(body: ResetReq, user=Depends(get_current_user)):
    """Wipe this user's data and seed exactly the strict v2.17 starter rows
    (Item 47): 2 tasks · 2 routines · 2 cash-flow entries. Other modules
    (notes/reminders/deadlines) start empty so the user makes them their own.
    """
    if body.confirm != "RESET":
        raise HTTPException(400, "Send {confirm:'RESET'} to proceed")
    uid = user["id"]
    for c in ("tasks", "routines", "transactions", "notes",
              "reminders", "deadlines", "routine_logs"):
        await db[c].delete_many({"user_id": uid})
    pid = await _ensure_default_project(uid)
    await _seed_strict_starter(uid, pid)
    return {"ok": True, "seeded": 6}


# ───────────────────── Reminders ─────────────────────
class ReminderIn(BaseModel):
    title: str
    notes: Optional[str] = ""
    fire_at: str  # ISO datetime in UTC
    recurrence: Optional[str] = "none"  # v2.2: free text — supports "every 15 days" etc.
    custom_recurrence: Optional[str] = None  # human-readable label for "custom" preset
    source_page: Optional[str] = None
    source_context: Optional[Dict[str, Any]] = None
    project_id: Optional[str] = None  # v2.17


class Reminder(BaseModel):
    id: str
    user_id: str
    title: str
    notes: str = ""
    fire_at: str
    recurrence: str = "none"
    custom_recurrence: Optional[str] = None
    sent: bool = False
    source_page: Optional[str] = None
    source_context: Optional[Dict[str, Any]] = None
    created_at: str


@api.get("/reminders", response_model=List[Reminder])
async def list_reminders(project_id: Optional[str] = None, user=Depends(get_current_user)):
    q: Dict[str, Any] = {"user_id": user["id"]}
    if project_id:
        q["project_id"] = project_id
    docs = await db.reminders.find(q, {"_id": 0}) \
        .sort("fire_at", 1).to_list(500)
    return [Reminder(**d) for d in docs]


@api.post("/reminders", response_model=Reminder)
async def create_reminder(body: ReminderIn, user=Depends(get_current_user)):
    project_id = body.project_id or await _ensure_default_project(user["id"])
    await _assert_project_write(user["id"], project_id)
    doc = {
        "id": new_id(), "user_id": user["id"],
        "title": body.title, "notes": body.notes or "",
        "fire_at": body.fire_at, "recurrence": body.recurrence or "none",
        "custom_recurrence": body.custom_recurrence,
        "source_page": body.source_page,
        "source_context": body.source_context,
        "project_id": project_id,
        "sent": False, "created_at": now_iso(),
    }
    await db.reminders.insert_one(dict(doc))
    return Reminder(**doc)


@api.patch("/reminders/{rid}", response_model=Reminder)
async def update_reminder(rid: str, body: Dict[str, Any], user=Depends(get_current_user)):
    allowed = {"title", "notes", "fire_at", "recurrence", "custom_recurrence",
               "sent", "source_page", "source_context", "project_id"}
    body = {k: v for k, v in body.items() if k in allowed}
    res = await db.reminders.find_one_and_update(
        {"id": rid, "user_id": user["id"]}, {"$set": body},
        projection={"_id": 0}, return_document=True,
    )
    if not res:
        raise HTTPException(404, "Reminder not found")
    return Reminder(**res)


@api.delete("/reminders/{rid}")
async def delete_reminder(rid: str, user=Depends(get_current_user)):
    await db.reminders.delete_one({"id": rid, "user_id": user["id"]})
    return {"ok": True}


@api.get("/reminders/{rid}/ics")
async def reminder_ics(rid: str, user=Depends(get_current_user)):
    r = await db.reminders.find_one({"id": rid, "user_id": user["id"]}, {"_id": 0})
    if not r:
        raise HTTPException(404, "Not found")
    dt = datetime.fromisoformat(r["fire_at"].replace("Z", "+00:00"))
    dt_utc = dt.astimezone(timezone.utc)
    stamp = dt_utc.strftime("%Y%m%dT%H%M%SZ")
    end = (dt_utc + timedelta(minutes=15)).strftime("%Y%m%dT%H%M%SZ")
    ics = (
        "BEGIN:VCALENDAR\r\nVERSION:2.0\r\nPRODID:-//Mind Matters//EN\r\n"
        "CALSCALE:GREGORIAN\r\nMETHOD:PUBLISH\r\n"
        "BEGIN:VEVENT\r\n"
        f"UID:{r['id']}@mindmatters\r\n"
        f"DTSTAMP:{stamp}\r\n"
        f"DTSTART:{stamp}\r\nDTEND:{end}\r\n"
        f"SUMMARY:{r['title']}\r\n"
        f"DESCRIPTION:{(r.get('notes') or '').replace(chr(10), ' ')}\r\n"
        "BEGIN:VALARM\r\nTRIGGER:-PT0M\r\nACTION:DISPLAY\r\n"
        f"DESCRIPTION:{r['title']}\r\nEND:VALARM\r\n"
        "END:VEVENT\r\nEND:VCALENDAR\r\n"
    )
    from fastapi.responses import Response
    return Response(
        content=ics.encode(),
        media_type="text/calendar",
        headers={"Content-Disposition": f'attachment; filename="reminder.ics"'},
    )


# ───────────────────── iCal subscription feed ─────────────────────
def _to_ics_dt(iso_str: str) -> str:
    try:
        dt = datetime.fromisoformat(iso_str.replace("Z", "+00:00"))
    except Exception:
        return ""
    if dt.tzinfo is None:
        dt = dt.replace(tzinfo=timezone.utc)
    return dt.astimezone(timezone.utc).strftime("%Y%m%dT%H%M%SZ")


def _rrule_for(recurrence: str) -> str:
    m = {
        "daily": "RRULE:FREQ=DAILY", "weekly": "RRULE:FREQ=WEEKLY",
        "monthly": "RRULE:FREQ=MONTHLY", "quarterly": "RRULE:FREQ=MONTHLY;INTERVAL=3",
        "half-yearly": "RRULE:FREQ=MONTHLY;INTERVAL=6", "yearly": "RRULE:FREQ=YEARLY",
    }
    return m.get((recurrence or "none").lower(), "")


@api.get("/cal/feed/token")
async def cal_feed_token_get(user=Depends(get_current_user)):
    u = await db.users.find_one({"id": user["id"]}, {"_id": 0})
    return {"token": (u or {}).get("cal_token")}


@api.post("/cal/feed/token")
async def cal_feed_token_rotate(user=Depends(get_current_user)):
    tok = secrets.token_urlsafe(24)
    await db.users.update_one({"id": user["id"]}, {"$set": {"cal_token": tok}})
    return {"token": tok}


@app.get("/api/cal/{token}.ics")
async def cal_feed_ics(token: str):
    """PUBLIC endpoint — subscribed from iOS/Google Calendar without auth.
    Token is long-random, rotatable, and the only identifier."""
    user = await db.users.find_one({"cal_token": token}, {"_id": 0})
    if not user:
        raise HTTPException(404, "Not found")
    reminders = await db.reminders.find({"user_id": user["id"]}, {"_id": 0}).to_list(2000)
    now_stamp = _to_ics_dt(now_iso())
    lines = ["BEGIN:VCALENDAR", "VERSION:2.0", "PRODID:-//Mind Matters//Reminders//EN",
             "CALSCALE:GREGORIAN", "METHOD:PUBLISH",
             f"X-WR-CALNAME:Mind Matters — {user.get('first_name','You')}"]
    for r in reminders:
        start = _to_ics_dt(r.get("fire_at", ""))
        if not start:
            continue
        uid = f"{r['id']}@mindmatters"
        title = (r.get("title") or "Reminder").replace("\n", " ")
        notes = (r.get("notes") or "").replace("\n", "\\n")
        lines += ["BEGIN:VEVENT", f"UID:{uid}", f"DTSTAMP:{now_stamp}",
                  f"DTSTART:{start}", f"DTEND:{start}", f"SUMMARY:{title}"]
        if notes:
            lines.append(f"DESCRIPTION:{notes}")
        rrule = _rrule_for(r.get("recurrence", "none"))
        if rrule:
            lines.append(rrule)
        lines += ["BEGIN:VALARM", "ACTION:DISPLAY", "DESCRIPTION:Reminder",
                  "TRIGGER:-PT10M", "END:VALARM", "END:VEVENT"]
    lines.append("END:VCALENDAR")
    body = "\r\n".join(lines) + "\r\n"
    from fastapi.responses import Response
    return Response(content=body, media_type="text/calendar; charset=utf-8",
                    headers={"Content-Disposition": f"inline; filename=mind-matters-{token[:8]}.ics"})


class ResendReq(BaseModel):
    fire_at: Optional[str] = None


@api.post("/reminders/{rid}/resend", response_model=Reminder)
async def reminder_resend(rid: str, body: ResendReq, user=Depends(get_current_user)):
    orig = await db.reminders.find_one({"id": rid, "user_id": user["id"]}, {"_id": 0})
    if not orig:
        raise HTTPException(404, "Reminder not found")
    fire_at = body.fire_at
    if not fire_at:
        rec = (orig.get("recurrence") or "none").lower()
        try:
            base = datetime.fromisoformat(orig["fire_at"].replace("Z", "+00:00"))
        except Exception:
            base = datetime.now(timezone.utc)
        delta = {"daily": 1, "weekly": 7, "monthly": 30, "quarterly": 90,
                 "half-yearly": 183, "yearly": 365}.get(rec, 1)
        fire_at = (base + timedelta(days=delta)).astimezone(timezone.utc).isoformat()
    doc = {
        "id": new_id(), "user_id": user["id"],
        "title": orig.get("title", "Reminder"), "notes": orig.get("notes", ""),
        "fire_at": fire_at, "recurrence": orig.get("recurrence", "none"),
        "source_page": orig.get("source_page"),
        "source_context": orig.get("source_context"),
        "sent": False, "created_at": now_iso(),
    }
    await db.reminders.insert_one(dict(doc))
    return Reminder(**doc)


# ───────────────────── All-data export (.xlsx) ─────────────────────
@api.get("/export/data.xlsx")
async def export_all_xlsx(user=Depends(get_current_user)):
    try:
        from openpyxl import Workbook
    except Exception as e:
        raise HTTPException(500, f"openpyxl not installed: {e}")
    wb = Workbook()
    wb.remove(wb.active)
    sources = {
        "Tasks": ("tasks", ["sr_no", "date", "group", "name", "task", "details", "status", "created_at"]),
        "Routines": ("routines", ["sr_no", "group", "name", "activity", "details", "frequency", "priority", "status", "created_at"]),
        "CashFlow": ("transactions", ["sr_no", "date", "group", "name", "details", "amount", "remarks", "head", "category", "mode", "created_at"]),
        "Notes": ("notes", ["title", "body", "tags", "pinned", "created_at"]),
        "Reminders": ("reminders", ["title", "notes", "fire_at", "recurrence", "source_page", "sent", "created_at"]),
        "Deadlines": ("deadlines", ["title", "due_date", "notes", "created_at"]),
    }
    for sheet_name, (coll_name, cols) in sources.items():
        rows = await db[coll_name].find({"user_id": user["id"]}, {"_id": 0}).to_list(5000)
        ws = wb.create_sheet(sheet_name)
        ws.append(cols)
        for r in rows:
            ws.append([str(r.get(c, "")) if r.get(c) is not None else "" for c in cols])
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    from fastapi.responses import Response
    return Response(
        content=buf.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": 'attachment; filename="mind-matters-export.xlsx"'},
    )


@api.get("/cashflow/upcoming-payments")
async def cashflow_upcoming_payments(user=Depends(get_current_user)):
    """Return reminders that originated on the cash-flow page plus any recurring
    transactions whose next fire date falls in the current calendar month.
    """
    today = date.today()
    month_start = today.replace(day=1)
    next_month = (month_start + timedelta(days=32)).replace(day=1)
    items = []
    # Reminders sourced from cash-flow
    async for r in db.reminders.find(
        {"user_id": user["id"], "source_page": {"$in": ["cash-flow", "cashflow"]},
         "sent": False}, {"_id": 0},
    ):
        try:
            fire_d = datetime.fromisoformat(r["fire_at"].replace("Z", "+00:00")).date()
        except Exception:
            continue
        if not (month_start <= fire_d < next_month):
            continue
        # Try to look up the linked transaction amount
        amt = 0.0
        if r.get("source_id"):
            tx = await db.transactions.find_one(
                {"id": r["source_id"], "user_id": user["id"]}, {"_id": 0},
            )
            if tx:
                amt = float(tx.get("amount") or 0)
        items.append({
            "title": r.get("title") or "Payment",
            "due_date": fire_d.isoformat(),
            "amount": amt,
            "reminder_id": r.get("id"),
            "source_id": r.get("source_id"),
        })
    items.sort(key=lambda x: x["due_date"])
    total = sum(it["amount"] for it in items)
    return {"items": items, "total": round(total), "month": month_start.strftime("%B %Y")}


@api.get("/calendar/feed.ics")
async def calendar_ics_feed(token: str = Query(""), user_id: str = Query("")):
    """iCal subscription feed of all of a user's tasks (with date) + reminders.
    Auth: pass ?token=<jwt> as a query param (URL must work in calendar apps
    that don't send Authorization headers).
    """
    user = None
    if token:
        try:
            payload = pyjwt.decode(token, JWT_SECRET, algorithms=["HS256"])
            user = await db.users.find_one({"id": payload.get("user_id") or payload.get("sub")}, {"_id": 0})
        except Exception:
            user = None
    if not user:
        raise HTTPException(401, "Token required")
    lines = [
        "BEGIN:VCALENDAR",
        "VERSION:2.0",
        "PRODID:-//Mind Matters//EN",
        "CALSCALE:GREGORIAN",
        "NAME:Mind Matters",
        "X-WR-CALNAME:Mind Matters",
    ]

    def _esc(s):
        return (s or "").replace("\\", "\\\\").replace(",", "\\,").replace(";", "\\;").replace("\n", "\\n")

    def _add_event(uid, summary, start_date, description=""):
        if not start_date:
            return
        try:
            d = start_date if isinstance(start_date, str) else start_date.isoformat()
            dt = d.replace("-", "")[:8]
        except Exception:
            return
        lines.extend([
            "BEGIN:VEVENT",
            f"UID:{uid}@mindmatters",
            f"DTSTAMP:{datetime.now(timezone.utc).strftime('%Y%m%dT%H%M%SZ')}",
            f"DTSTART;VALUE=DATE:{dt}",
            f"SUMMARY:{_esc(summary)}",
            f"DESCRIPTION:{_esc(description)}",
            "END:VEVENT",
        ])

    async for t in db.tasks.find({"user_id": user["id"]}, {"_id": 0}):
        if t.get("date"):
            _add_event(t["id"], t.get("task") or "Task", t["date"],
                       f"{t.get('group') or ''} · {t.get('name') or ''}")
    async for r in db.reminders.find({"user_id": user["id"]}, {"_id": 0}):
        fa = r.get("fire_at") or ""
        if fa:
            try:
                d = datetime.fromisoformat(fa.replace("Z", "+00:00")).date().isoformat()
            except Exception:
                d = None
            if d:
                _add_event(r["id"], r.get("title") or "Reminder", d, r.get("notes") or "")
    lines.append("END:VCALENDAR")
    body = "\r\n".join(lines)
    from fastapi.responses import Response
    return Response(content=body, media_type="text/calendar; charset=utf-8")


# ───────────────────── Universal attachments (tasks/routines/transactions/notes) ─────────────────────
_ATTACH_COLLECTIONS = {
    "tasks": "tasks",
    "routines": "routines",
    "transactions": "transactions",
    "notes": "notes",
}


@api.post("/{module}/{rid}/attachments")
async def upload_row_attachment(
    module: str,
    rid: str,
    file: UploadFile = File(...),
    user=Depends(get_current_user),
):
    coll = _ATTACH_COLLECTIONS.get(module)
    if not coll:
        raise HTTPException(404, "Unknown module")
    row = await db[coll].find_one({"id": rid, "user_id": user["id"]}, {"_id": 0})
    if not row:
        raise HTTPException(404, "Row not found")
    content = await file.read()
    if len(content) > 10 * 1024 * 1024:
        raise HTTPException(413, "Attachment too large (max 10MB)")
    current = row.get("attachments") or []
    if len(current) >= 10:
        raise HTTPException(413, "Max 10 attachments per row")
    mime = file.content_type or "application/octet-stream"
    data_url = f"data:{mime};base64,{base64.b64encode(content).decode()}"
    entry = {
        "id": new_id()[:10],
        "name": file.filename or "file",
        "mime": mime,
        "size": len(content),
        "data_url": data_url,
    }
    current.append(entry)
    await db[coll].update_one(
        {"id": rid, "user_id": user["id"]},
        {"$set": {"attachments": current, "updated_at": now_iso()}},
    )
    return entry


@api.delete("/{module}/{rid}/attachments/{att_id}")
async def delete_row_attachment(module: str, rid: str, att_id: str, user=Depends(get_current_user)):
    coll = _ATTACH_COLLECTIONS.get(module)
    if not coll:
        raise HTTPException(404, "Unknown module")
    row = await db[coll].find_one({"id": rid, "user_id": user["id"]}, {"_id": 0})
    if not row:
        raise HTTPException(404, "Row not found")
    remaining = [a for a in (row.get("attachments") or []) if a.get("id") != att_id]
    await db[coll].update_one(
        {"id": rid, "user_id": user["id"]},
        {"$set": {"attachments": remaining, "updated_at": now_iso()}},
    )
    return {"ok": True}


@api.post("/seed/first-login")
async def seed_first_login(user=Depends(get_current_user)):
    """Insert the strict v2.17 starter rows (Item 47) ONLY if all three
    starter collections are empty for this user. Idempotent — safe to
    call on every login. Returns counts seeded.
    """
    tasks_cnt = await db.tasks.count_documents({"user_id": user["id"]})
    routines_cnt = await db.routines.count_documents({"user_id": user["id"]})
    tx_cnt = await db.transactions.count_documents({"user_id": user["id"]})
    if tasks_cnt or routines_cnt or tx_cnt:
        return {"seeded": False, "reason": "already has data"}
    pid = await _ensure_default_project(user["id"])
    await _seed_strict_starter(user["id"], pid)
    return {"seeded": True, "tasks": 2, "routines": 2, "transactions": 2}


async def _seed_strict_starter(uid: str, project_id: str) -> None:
    """Insert the EXACT strict starter rows specified by Item 47:
    - Tasks: 2 (Rahul Courier, Amit Invoice follow-up)
    - Routines: 2 (Uptime, Hydrate & Tea)
    - Cash Flow: 2 (Zomato expense, Brinda loan)
    Notes / reminders / deadlines stay empty for the user to fill in.
    """
    today_iso = date.today().isoformat()
    now = now_iso()
    common = {"user_id": uid, "created_at": now, "updated_at": now,
              "project_id": project_id}
    await db.tasks.insert_many([
        {"id": new_id(), "sr_no": 1, "order_index": 0,
         "date": today_iso, "group": "Work", "section": "",
         "name": "Rahul", "task": "Courier",
         "details": "Partnership Agreement Copy To Karan M",
         "status": "Delegate", "parent_id": None, "flagged": False,
         "attachments": [], **common},
        {"id": new_id(), "sr_no": 2, "order_index": 1,
         "date": today_iso, "group": "Work", "section": "",
         "name": "Amit", "task": "Invoice follow-up",
         "details": "Check if Q2 invoice was received",
         "status": "Follow-Up", "parent_id": None, "flagged": False,
         "attachments": [], **common},
    ])
    await db.routines.insert_many([
        {"id": new_id(), "sr_no": 1, "order_index": 0,
         "group": "3 Hours", "name": "Karan", "activity": "Uptime",
         "details": "Rise & Shine", "frequency": "Daily",
         "priority": "Medium", "status": "Active", "section": "",
         "parent_id": None, "flagged": False, "attachments": [], **common},
        {"id": new_id(), "sr_no": 2, "order_index": 1,
         "group": "3 Hours", "name": "Karan", "activity": "Hydrate & Tea",
         "details": "Health routine", "frequency": "Daily",
         "priority": "Medium", "status": "Active", "section": "",
         "parent_id": None, "flagged": False, "attachments": [], **common},
    ])
    await db.transactions.insert_many([
        {"id": new_id(), "sr_no": 1, "order_index": 0,
         "date": today_iso, "amount": 555.0, "name": "Zomato", "vendor": "Zomato",
         "details": "Invoice number 3", "mode": "UPI",
         "head": "Food & Beverages", "category": "expense",
         "group": "Company 1", "section": "",
         "expense_head": "Food & Beverages", "direction": "out",
         "account": "Personal", "notes": "", "currency": "INR",
         "parent_id": None, "flagged": False, "attachments": [],
         "source": "manual", **common},
        {"id": new_id(), "sr_no": 2, "order_index": 1,
         "date": today_iso, "amount": 50000.0, "name": "Brinda", "vendor": "Brinda",
         "details": "Lent at 9% pa", "mode": "Bank",
         "head": "Personal Loan", "category": "loan_given",
         "group": "Personal", "section": "",
         "expense_head": "Personal Loan", "direction": "out",
         "account": "Personal", "notes": "",
         "interest_rate": 9.0, "currency": "INR",
         "parent_id": None, "flagged": False, "attachments": [],
         "source": "manual", **common},
    ])


@api.post("/admin/wipe-all-data")
async def admin_wipe_all_data(user=Depends(get_current_user)):
    """Dev/maintenance: wipe ALL data rows for the current user (keeps account)."""
    for coll in ("tasks", "routines", "transactions", "notes", "reminders",
                 "deadlines", "loans", "investments", "affirmations", "routine_logs"):
        await db[coll].delete_many({"user_id": user["id"]})
    return {"ok": True, "user_id": user["id"]}


@api.get("/news")
async def news_today(category: str = "all", user=Depends(get_current_user)):
    """Lightweight news widget — fetches headlines from a free RSS source per category.

    Categories: all | business | tech | india | world. Falls back to a stable
    static list if the network is unavailable so the dashboard never feels broken.
    """
    feeds = {
        "all": "https://news.google.com/rss?hl=en-IN&gl=IN&ceid=IN:en",
        "business": "https://news.google.com/rss/headlines/section/topic/BUSINESS?hl=en-IN&gl=IN&ceid=IN:en",
        "tech": "https://news.google.com/rss/headlines/section/topic/TECHNOLOGY?hl=en-IN&gl=IN&ceid=IN:en",
        "india": "https://news.google.com/rss/headlines/section/topic/NATION?hl=en-IN&gl=IN&ceid=IN:en",
        "world": "https://news.google.com/rss/headlines/section/topic/WORLD?hl=en-IN&gl=IN&ceid=IN:en",
    }
    url = feeds.get(category, feeds["all"])
    try:
        import urllib.request
        import re as _re
        req = urllib.request.Request(url, headers={"User-Agent": "Mind Matters/1.0"})
        with urllib.request.urlopen(req, timeout=4) as resp:
            xml = resp.read().decode("utf-8", "ignore")
        items = _re.findall(r"<item>(.*?)</item>", xml, flags=_re.DOTALL)[:5]
        out = []
        for it in items:
            tm = _re.search(r"<title>(.*?)</title>", it, flags=_re.DOTALL)
            lm = _re.search(r"<link>(.*?)</link>", it, flags=_re.DOTALL)
            title = (tm.group(1) if tm else "").strip()
            if title.startswith("<![CDATA["):
                title = title[9:].rstrip("]>").strip()
            out.append({"title": title, "url": (lm.group(1).strip() if lm else "")})
        return {"category": category, "items": out, "source": "google-news"}
    except Exception:
        return {"category": category, "items": [
            {"title": "News unavailable — check your connection", "url": ""},
        ], "source": "fallback"}


# ───────────────────── Currency conversion (cash-flow summary tiles) ─────────────────────
_FX_CACHE: Dict[str, Any] = {"rates": {}, "fetched_at": 0}


async def _fetch_fx_rates(base: str = "INR") -> Dict[str, float]:
    """Return a {currency_code: rate_to_base} map. Cache for 6 hours.

    Uses exchangerate.host (free, no key). Falls back to a static table when
    the network call fails so summary tiles never go blank.
    """
    base = (base or "INR").upper()
    now = time.time()
    if (now - _FX_CACHE["fetched_at"]) < 6 * 3600 and base in _FX_CACHE["rates"]:
        return _FX_CACHE["rates"][base]
    fallback = {  # ~Feb 2026 indicative rates to INR
        "INR": 1.0, "USD": 86.0, "EUR": 94.0, "GBP": 110.0,
        "JPY": 0.57, "AED": 23.4, "CAD": 60.5, "AUD": 56.0,
    }
    try:
        import urllib.request
        url = f"https://api.exchangerate.host/latest?base={base}"
        req = urllib.request.Request(url, headers={"User-Agent": "Mind Matters/1.0"})
        with urllib.request.urlopen(req, timeout=4) as resp:
            data = json.loads(resp.read().decode("utf-8"))
        # exchangerate.host returns rates FROM base, so to convert FROM `code` to
        # `base` we invert: rate_to_base = 1 / data.rates[code].
        rates = {base: 1.0}
        for code, r in (data.get("rates") or {}).items():
            try:
                rates[code] = 1.0 / float(r) if r else 0
            except Exception:
                pass
        if len(rates) < 2:
            raise ValueError("empty fx response")
        _FX_CACHE["rates"][base] = rates
        _FX_CACHE["fetched_at"] = now
        return rates
    except Exception:
        # Translate fallback (INR-based) → arbitrary base
        if base == "INR":
            return fallback
        b = fallback.get(base, 1.0) or 1.0
        return {c: (v / b) for c, v in fallback.items()}


@api.get("/fx/rates")
async def fx_rates(base: str = "INR", user=Depends(get_current_user)):
    rates = await _fetch_fx_rates(base)
    return {"base": base.upper(), "rates": rates, "fetched_at": _FX_CACHE["fetched_at"]}


@api.get("/cashflow/totals")
async def cashflow_totals(base: str = "INR", user=Depends(get_current_user)):
    """Aggregate transactions into per-category totals converted to `base`."""
    rates = await _fetch_fx_rates(base.upper())
    out = {"income": 0.0, "expense": 0.0, "asset": 0.0, "liability": 0.0,
           "loan_given": 0.0, "loan_taken": 0.0}
    async for t in db.transactions.find({"user_id": user["id"]}, {"_id": 0}):
        cat = t.get("category") or ("income" if t.get("direction") == "in" else "expense")
        if cat not in out:
            out[cat] = 0.0
        try:
            amt = float(t.get("amount") or 0)
            cur = (t.get("currency") or "INR").upper()
            rate = rates.get(cur, 1.0)
            out[cat] += amt * rate
        except Exception:
            pass
    return {"base": base.upper(), "totals": {k: round(v) for k, v in out.items()}}


# ───────────────────── Reports + Calendar + Patterns ─────────────────────
@api.get("/reports/timeline")
async def reports_timeline(days: int = 30, user=Depends(get_current_user)):
    """Return a chronological feed of recent activity across modules."""
    since = (datetime.now(timezone.utc) - timedelta(days=days)).isoformat()
    items = []
    async for t in db.tasks.find(
        {"user_id": user["id"], "created_at": {"$gte": since}}, {"_id": 0},
    ).sort("created_at", -1).limit(120):
        items.append({
            "ts": t.get("created_at"),
            "kind": "task",
            "title": t.get("task") or "(task)",
            "subtitle": " · ".join(
                [b for b in [t.get("group"), t.get("name"), t.get("status")] if b]
            ),
        })
    async for t in db.transactions.find(
        {"user_id": user["id"], "created_at": {"$gte": since}}, {"_id": 0},
    ).sort("created_at", -1).limit(120):
        try:
            amt = f"₹{float(t.get('amount') or 0):,.0f}"
        except Exception:
            amt = str(t.get("amount") or "")
        items.append({
            "ts": t.get("created_at"),
            "kind": "transaction",
            "title": f"{amt} · {t.get('vendor') or t.get('details') or 'expense'}",
            "subtitle": " · ".join(
                [b for b in [t.get("category"), t.get("head")] if b]
            ),
        })
    async for n in db.notes.find(
        {"user_id": user["id"], "created_at": {"$gte": since}}, {"_id": 0},
    ).sort("created_at", -1).limit(80):
        items.append({
            "ts": n.get("created_at"),
            "kind": "note",
            "title": n.get("title") or "(note)",
            "subtitle": ", ".join(n.get("tags") or []),
        })
    items.sort(key=lambda x: x.get("ts") or "", reverse=True)
    return items[:200]


@api.get("/reports/cashflow-monthly")
async def reports_cashflow_monthly(months: int = 6, user=Depends(get_current_user)):
    """Aggregate transactions into monthly totals by category."""
    today = date.today()
    start = (today.replace(day=1) - timedelta(days=31 * (months - 1))).replace(day=1)
    buckets: Dict[str, Dict[str, float]] = {}
    async for t in db.transactions.find({"user_id": user["id"]}, {"_id": 0}):
        d = t.get("date") or (t.get("created_at") or "")[:10]
        if not d or len(d) < 7:
            continue
        yyyy_mm = d[:7]
        if yyyy_mm < start.isoformat()[:7]:
            continue
        cat = t.get("category") or ("income" if t.get("direction") == "in" else "expense")
        b = buckets.setdefault(yyyy_mm, {"income": 0, "expense": 0, "asset": 0, "liability": 0})
        try:
            b[cat] = b.get(cat, 0) + float(t.get("amount") or 0)
        except Exception:
            pass
    return [{"month": m, **v} for m, v in sorted(buckets.items())]


@api.get("/reports/patterns")
async def reports_patterns(user=Depends(get_current_user)):
    """Rule-based pattern radar — spending spikes, overdue tasks, upcoming loans, routine streaks."""
    patterns = []
    today = date.today()
    this_start = today.replace(day=1)
    last_end = this_start - timedelta(days=1)
    last_start = last_end.replace(day=1)
    this_exp = 0.0
    last_exp = 0.0
    async for t in db.transactions.find(
        {"user_id": user["id"], "category": "expense"}, {"_id": 0},
    ):
        d = t.get("date")
        if not d:
            continue
        try:
            dd = date.fromisoformat(d)
            amt = float(t.get("amount") or 0)
            if dd >= this_start:
                this_exp += amt
            elif last_start <= dd <= last_end:
                last_exp += amt
        except Exception:
            pass
    if last_exp > 0:
        delta_pct = ((this_exp - last_exp) / last_exp) * 100
        if delta_pct >= 20:
            patterns.append({
                "severity": "alert",
                "title": f"Spending up {delta_pct:.0f}% vs last month",
                "detail": f"This month ₹{this_exp:,.0f} vs last month ₹{last_exp:,.0f}",
            })
        elif delta_pct <= -20:
            patterns.append({
                "severity": "info",
                "title": f"Spending down {abs(delta_pct):.0f}% vs last month",
                "detail": f"This month ₹{this_exp:,.0f} vs last month ₹{last_exp:,.0f} — keep it up.",
            })
    overdue = 0
    async for t in db.tasks.find(
        {"user_id": user["id"], "status": {"$ne": "Completed"}}, {"_id": 0},
    ):
        d = t.get("date")
        if not d:
            continue
        try:
            if date.fromisoformat(d) < today:
                overdue += 1
        except Exception:
            pass
    if overdue:
        patterns.append({
            "severity": "warn",
            "title": f"{overdue} overdue task{'s' if overdue != 1 else ''}",
            "detail": "Tasks past their date with status still open.",
        })
    upcoming_loans = []
    async for t in db.transactions.find(
        {"user_id": user["id"], "category": {"$in": ["liability", "asset"]}}, {"_id": 0},
    ):
        rep = t.get("repayment_date")
        if not rep:
            continue
        try:
            rd = date.fromisoformat(rep)
            days = (rd - today).days
            if 0 <= days <= 14:
                upcoming_loans.append((days, t))
        except Exception:
            pass
    if upcoming_loans:
        upcoming_loans.sort()
        days, t = upcoming_loans[0]
        patterns.append({
            "severity": "warn" if days <= 3 else "info",
            "title": f"Loan repayment in {days} day{'s' if days != 1 else ''}",
            "detail": f"{t.get('vendor') or t.get('name') or 'Loan'} · ₹{float(t.get('amount') or 0):,.0f}",
        })
    return patterns


@api.get("/reports/ai-patterns")
async def reports_ai_patterns(user=Depends(get_current_user)):
    """Ask Gemini to surface non-obvious patterns in the user's last 60 days.

    Returns a list of {title, detail} strings. Falls back to an empty list when
    the LLM key is missing or the call fails — the rule-based /patterns
    endpoint always succeeds, so the UI degrades gracefully.
    """
    if not EMERGENT_LLM_KEY:
        return []
    today = date.today()
    cutoff = today - timedelta(days=60)
    # Build a compact, anonymised snapshot for the LLM
    expenses = []
    async for t in db.transactions.find(
        {"user_id": user["id"], "category": "expense"}, {"_id": 0},
    ).sort("date", -1).limit(120):
        d = t.get("date")
        if not d:
            continue
        try:
            if date.fromisoformat(d) < cutoff:
                continue
        except Exception:
            continue
        expenses.append({
            "date": d,
            "amount": float(t.get("amount") or 0),
            "head": t.get("head") or t.get("expense_head") or "Uncategorized",
            "vendor": t.get("vendor") or t.get("name") or "",
        })
    tasks_summary = []
    async for t in db.tasks.find(
        {"user_id": user["id"]}, {"_id": 0},
    ).sort("updated_at", -1).limit(60):
        tasks_summary.append({
            "date": t.get("date"),
            "status": t.get("status"),
            "group": t.get("group"),
        })
    routine_logs = []
    async for r in db.routine_logs.find(
        {"user_id": user["id"], "date": {"$gte": cutoff.isoformat()}}, {"_id": 0},
    ).limit(200):
        routine_logs.append({
            "date": r.get("date"),
            "done": r.get("done"),
            "weekday": (date.fromisoformat(r["date"]).strftime("%A")
                        if r.get("date") else None),
        })
    snapshot = {
        "today": today.isoformat(),
        "expenses_last_60d": expenses[:60],
        "tasks_recent": tasks_summary[:40],
        "routine_logs_last_60d": routine_logs,
    }
    system = (
        "You are a pragmatic personal-analytics coach. Look at the user's "
        "last 60 days and surface 3-5 NON-OBVIOUS patterns a rule-based system "
        "would miss — e.g. weekday-specific habits, vendor concentration, "
        "category creep, recurring weekly spend, missed routine cadence. "
        "Each insight must be ACTIONABLE and SPECIFIC (mention day-of-week, "
        "vendor, or category whenever possible). Return STRICT JSON: a list "
        "of objects {title, detail}. No prose, no markdown, no preamble."
    )
    chat = LlmChat(api_key=EMERGENT_LLM_KEY, session_id=f"aip-{new_id()}",
                   system_message=system).with_model("gemini", "gemini-3-flash-preview")
    try:
        resp = await chat.send_message(UserMessage(
            text=json.dumps(snapshot, default=str)[:12000],
        ))
        t = (resp or "").strip()
        if t.startswith("```"):
            t = t.strip("`")
            if t.startswith("json"):
                t = t[4:]
        arr = json.loads(t)
        if not isinstance(arr, list):
            return []
        # Sanitize each item
        out = []
        for it in arr[:6]:
            if not isinstance(it, dict):
                continue
            title = str(it.get("title") or "").strip()
            detail = str(it.get("detail") or "").strip()
            if title:
                out.append({"title": title[:140], "detail": detail[:240]})
        return out
    except Exception:
        return []


@api.post("/reports/briefing")
async def reports_briefing(user=Depends(get_current_user)):
    """AI-generated weekly briefing."""
    today = date.today()
    week_ago = today - timedelta(days=7)
    tasks_done = await db.tasks.count_documents(
        {"user_id": user["id"], "status": "Completed",
         "updated_at": {"$gte": week_ago.isoformat()}},
    )
    tasks_pending = await db.tasks.count_documents(
        {"user_id": user["id"], "status": {"$ne": "Completed"}},
    )
    week_exp = 0.0
    async for t in db.transactions.find(
        {"user_id": user["id"], "category": "expense"}, {"_id": 0},
    ):
        try:
            d = t.get("date")
            if d and date.fromisoformat(d) >= week_ago:
                week_exp += float(t.get("amount") or 0)
        except Exception:
            pass
    patterns = await reports_patterns(user=user)
    snapshot = {
        "tasks_completed_this_week": tasks_done,
        "tasks_open": tasks_pending,
        "expense_this_week": round(week_exp),
        "patterns": [p["title"] for p in patterns],
    }
    fallback = (
        f"This week you ticked off {tasks_done} task{'s' if tasks_done != 1 else ''} "
        f"with {tasks_pending} still open. You spent about ₹{week_exp:,.0f}. "
        + (("Heads up: " + "; ".join(p["title"] for p in patterns) + ".") if patterns else "")
    ).strip()
    if not EMERGENT_LLM_KEY:
        return {"summary": fallback, "snapshot": snapshot}
    system = (
        "You are a calm, pragmatic life coach. Write a 3-4 sentence briefing for the user "
        "based on their last week. Warm but direct. Mention one win and one concrete next step. "
        "No emojis, no markdown, no preamble."
    )
    chat = LlmChat(api_key=EMERGENT_LLM_KEY, session_id=f"brief-{new_id()}",
                   system_message=system).with_model("gemini", "gemini-3-flash-preview")
    try:
        resp = await chat.send_message(UserMessage(
            text=f"Snapshot: {json.dumps(snapshot, default=str)}",
        ))
        return {"summary": (resp or "").strip() or fallback, "snapshot": snapshot}
    except Exception:
        return {"summary": fallback, "snapshot": snapshot}


# ───────────────────── Per-module exports (CSV + PDF) ─────────────────────
@api.get("/cashflow/loan-summary")
async def cashflow_loan_summary(user=Depends(get_current_user)):
    """Aggregate liability + asset(loan) rows for the dashboard widget.
    Returns: total_monthly_emi, active_loans, next_repayment (or None).
    A loan is considered "active" if it has interest_rate OR emi OR
    repayment_date set, and (when repayment_date is set) hasn't passed yet.
    """
    cur = db.transactions.find(
        {"user_id": user["id"], "category": {"$in": ["liability", "asset"]}},
        {"_id": 0},
    )
    docs = await cur.to_list(2000)
    today = date.today()
    total_emi = 0.0
    active = []
    next_rep = None
    for d in docs:
        rate = d.get("interest_rate")
        emi = d.get("emi")
        rep = d.get("repayment_date")
        if not (rate or emi or rep):
            continue
        # Active if no repayment date OR repayment in the future
        rep_d = None
        if rep:
            try:
                rep_d = date.fromisoformat(rep)
            except Exception:
                rep_d = None
        if rep_d and rep_d < today:
            continue
        active.append(d)
        # Compute monthly EMI if not stored
        if emi:
            try:
                total_emi += float(emi)
            except Exception:
                pass
        elif rep_d and d.get("date") and rate is not None:
            try:
                start_d = date.fromisoformat(d["date"])
                months = (rep_d.year - start_d.year) * 12 + (rep_d.month - start_d.month)
                P = float(d.get("amount") or 0)
                r = float(rate) / 100 / 12 if rate else 0
                if months > 0 and P > 0:
                    if r:
                        pow_ = (1 + r) ** months
                        total_emi += (P * r * pow_) / (pow_ - 1)
                    else:
                        total_emi += P / months
            except Exception:
                pass
        # Track nearest upcoming repayment
        if rep_d:
            if next_rep is None or rep_d < date.fromisoformat(next_rep["repayment_date"]):
                next_rep = {
                    "repayment_date": rep,
                    "days_until": (rep_d - today).days,
                    "vendor": d.get("vendor") or d.get("name") or "",
                    "amount": d.get("amount") or 0,
                    "category": d.get("category"),
                }
    return {
        "total_monthly_emi": round(total_emi),
        "active_loans": len(active),
        "next_repayment": next_rep,
    }


def _csv_response(filename: str, headers: List[str], rows: List[List[Any]]):
    import csv as _csv
    buf = io.StringIO()
    w = _csv.writer(buf)
    w.writerow(headers)
    for r in rows:
        w.writerow(["" if v is None else str(v) for v in r])
    from fastapi.responses import Response
    return Response(
        content=buf.getvalue(),
        media_type="text/csv; charset=utf-8",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _pdf_response(filename: str, title: str, headers: List[str], rows: List[List[Any]]):
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib import colors
    from reportlab.lib.units import mm
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=landscape(A4),
                            leftMargin=12 * mm, rightMargin=12 * mm,
                            topMargin=12 * mm, bottomMargin=12 * mm)
    styles = getSampleStyleSheet()
    story = [Paragraph(f"<b>{title}</b>", styles["Title"]), Spacer(1, 6)]
    data = [headers] + [[("" if v is None else str(v))[:80] for v in r] for r in rows]
    t = Table(data, repeatRows=1)
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0A0A0A")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#C9A961")),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 7.5),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#C9A961")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.white]),
    ]))
    story.append(t)
    doc.build(story)
    buf.seek(0)
    from fastapi.responses import Response
    return Response(
        content=buf.getvalue(),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


_EXPORT_DEFS: Dict[str, Dict[str, Any]] = {
    "tasks": {
        "collection": "tasks",
        "title": "Tasks",
        "headers": ["Sr", "Date", "Group", "Section", "To", "Task", "Details", "Status"],
        "fields": ["sr_no", "date", "group", "section", "name", "task", "details", "status"],
    },
    "routines": {
        "collection": "routines",
        "title": "Routines",
        "headers": ["Sr", "Group", "Section", "Name", "Task", "Details", "Frequency"],
        "fields": ["sr_no", "group", "section", "name", "activity", "details", "frequency"],
    },
    "cashflow": {
        "collection": "transactions",
        "title": "Cash Flow",
        "headers": ["Sr", "Date", "Group", "Section", "Vendor", "Details", "Amount",
                    "Mode", "Head", "Category", "Interest %", "Repayment", "EMI"],
        "fields": ["sr_no", "date", "group", "section", "vendor", "details", "amount",
                   "mode", "head", "category", "interest_rate", "repayment_date", "emi"],
    },
    "notes": {
        "collection": "notes",
        "title": "Notes",
        "headers": ["Title", "Body", "Tags", "Pinned", "Created"],
        "fields": ["title", "body", "tags", "pinned", "created_at"],
    },
    "reminders": {
        "collection": "reminders",
        "title": "Reminders",
        "headers": ["Title", "Notes", "Fire at", "Recurrence", "Source", "Sent"],
        "fields": ["title", "notes", "fire_at", "recurrence", "source_page", "sent"],
    },
}


@api.get("/export/{module}.csv")
async def export_module_csv(module: str, user=Depends(get_current_user)):
    spec = _EXPORT_DEFS.get(module.lower())
    if not spec:
        raise HTTPException(404, "Unknown module")
    docs = await db[spec["collection"]].find({"user_id": user["id"]}, {"_id": 0}).to_list(10000)
    docs.sort(key=lambda d: (d.get("sr_no") or 0, d.get("created_at") or ""))
    rows = [[d.get(f, "") for f in spec["fields"]] for d in docs]
    return _csv_response(f"mind-matters-{module}.csv", spec["headers"], rows)


@api.get("/export/{module}.pdf")
async def export_module_pdf(module: str, user=Depends(get_current_user)):
    spec = _EXPORT_DEFS.get(module.lower())
    if not spec:
        raise HTTPException(404, "Unknown module")
    docs = await db[spec["collection"]].find({"user_id": user["id"]}, {"_id": 0}).to_list(10000)
    docs.sort(key=lambda d: (d.get("sr_no") or 0, d.get("created_at") or ""))
    rows = [[d.get(f, "") for f in spec["fields"]] for d in docs]
    return _pdf_response(f"mind-matters-{module}.pdf", spec["title"], spec["headers"], rows)


# ───────────────────── Change password ─────────────────────
class ChangePasswordReq(BaseModel):
    current_password: str
    new_password: str


@api.post("/auth/change-password")
async def change_password(body: ChangePasswordReq, user=Depends(get_current_user)):
    if len(body.new_password) < 6:
        raise HTTPException(400, "Password must be at least 6 characters")
    u = await db.users.find_one({"id": user["id"]}, {"_id": 0})
    if not u or not u.get("password_hash"):
        raise HTTPException(400, "No password set — use signup instead")
    if not _verify_password(body.current_password, u["password_hash"]):
        raise HTTPException(401, "Current password is wrong")
    await db.users.update_one(
        {"id": user["id"]},
        {"$set": {"password_hash": _hash_password(body.new_password)}},
    )
    return {"ok": True}


# ───────────────────── Telegram linking + test ─────────────────────
@api.get("/telegram/status")
async def tg_status(user=Depends(get_current_user)):
    from tg import TG_TOKEN
    me = user
    chat_id = me.get("telegram_chat_id")
    bot_username = None
    if TG_TOKEN:
        try:
            async with httpx.AsyncClient(timeout=6) as cli:
                r = await cli.get(f"https://api.telegram.org/bot{TG_TOKEN}/getMe")
                bot_username = (r.json().get("result") or {}).get("username")
        except Exception:
            pass
    return {"linked": bool(chat_id), "chat_id": chat_id,
            "bot_username": bot_username,
            "configured": bool(TG_TOKEN)}


@api.post("/telegram/link-code")
async def tg_link_code(user=Depends(get_current_user)):
    code = secrets.token_urlsafe(8).replace("_", "").replace("-", "")[:10]
    doc = {
        "id": new_id(), "code": code, "user_id": user["id"],
        "used": False, "created_at": now_iso(),
    }
    await db.tg_links.insert_one(dict(doc))
    return {"code": code}


@api.post("/telegram/unlink")
async def tg_unlink(user=Depends(get_current_user)):
    await db.users.update_one({"id": user["id"]}, {"$set": {"telegram_chat_id": None}})
    return {"ok": True}


class TgTestReq(BaseModel):
    text: str = "Mind Matters test ping ✓"


@api.post("/telegram/send-test")
async def tg_send_test(body: TgTestReq, user=Depends(get_current_user)):
    cid = user.get("telegram_chat_id")
    if not cid:
        raise HTTPException(400, "Telegram not linked")
    res = await tg_send(cid, body.text)
    return {"ok": bool(res and res.get("ok")), "response": res}


# ───────────────────── Share statements via Telegram ─────────────────────
class ShareStatementReq(BaseModel):
    kind: Literal["loan", "tasks", "transactions"] = "loan"
    name: Optional[str] = None  # filter by person (loan.name / task.name)
    month: Optional[str] = None  # YYYY-MM for transactions


@api.post("/share/statement")
async def share_statement(body: ShareStatementReq, user=Depends(get_current_user)):
    cid = user.get("telegram_chat_id")
    if not cid:
        raise HTTPException(400, "Telegram not linked. Connect in Settings.")

    if body.kind == "loan":
        q: Dict[str, Any] = {"user_id": user["id"]}
        if body.name:
            q["name"] = {"$regex": f"^{body.name}$", "$options": "i"}
        loans = await db.loans.find(q, {"_id": 0}).sort("date", 1).to_list(1000)
        if not loans:
            raise HTTPException(404, f"No loans found{f' for {body.name}' if body.name else ''}")
        rows = []
        total_g = total_t = total_i = 0
        for l in loans:
            acc = _accrued_interest(l["amount"], l.get("interest", 0), l.get("date"))
            rows.append([
                str(l.get("sr_no")), l.get("date", ""), l["name"], l["status"],
                f"₹ {l['amount']:,.0f}", f"{l.get('interest', 0)}%", f"₹ {acc:,.0f}",
                l.get("repayment_date") or "—",
            ])
            if l["status"] == "Given":
                total_g += l["amount"]
            if l["status"] == "Taken":
                total_t += l["amount"]
            total_i += acc
        pdf = render_simple_statement(
            title=f"Loan Statement{' — ' + body.name if body.name else ''}",
            subtitle=f"Prepared for {user['first_name']} · {datetime.now().strftime('%d %b %Y')}",
            meta={"User": user["first_name"], "Records": len(loans),
                  "Total Given": f"₹ {total_g:,.0f}", "Total Taken": f"₹ {total_t:,.0f}",
                  "Net Exposure": f"₹ {total_g - total_t:,.0f}",
                  "Interest Accrued": f"₹ {total_i:,.0f}"},
            table_headers=["Sr", "Date", "Name", "Status", "Amount", "Rate", "Accrued", "Repay by"],
            table_rows=rows,
            footer="Generated by Mind Matters — personal operating system.",
        )
        fname = f"loan-statement{('-' + body.name) if body.name else ''}.pdf"
        caption = f"Loan statement{' for ' + body.name if body.name else ''} · {len(loans)} record(s)"

    elif body.kind == "tasks":
        q = {"user_id": user["id"]}
        if body.name:
            q["name"] = {"$regex": f"^{body.name}$", "$options": "i"}
        tasks = await db.tasks.find(q, {"_id": 0}).sort("sr_no", 1).to_list(2000)
        rows = [[str(t.get("sr_no")), t.get("date", ""), t["name"] or "—", t["task"], t["status"]]
                for t in tasks]
        pdf = render_simple_statement(
            title=f"Tasks{' — ' + body.name if body.name else ''}",
            subtitle=f"{len(tasks)} task(s) · {datetime.now().strftime('%d %b %Y')}",
            meta={"Pending": sum(1 for t in tasks if t["status"] == "Pending"),
                  "Follow-Up": sum(1 for t in tasks if t["status"] == "Follow-Up"),
                  "Done": sum(1 for t in tasks if t["status"] == "Done")},
            table_headers=["Sr", "Date", "Person", "Task", "Status"],
            table_rows=rows,
            footer="Generated by Mind Matters.",
        )
        fname = f"tasks{('-' + body.name) if body.name else ''}.pdf"
        caption = f"Task statement{' for ' + body.name if body.name else ''}"

    else:  # transactions
        month = body.month or datetime.now(timezone.utc).strftime("%Y-%m")
        tx = await db.transactions.find(
            {"user_id": user["id"], "date": {"$regex": f"^{month}"}}, {"_id": 0}
        ).sort("date", 1).to_list(5000)
        rows = [[t.get("date", ""), t.get("direction"), f"₹ {t['amount']:,.0f}",
                 t.get("expense_head", ""), t.get("company", ""), t.get("mode", "")] for t in tx]
        total_out = sum(t["amount"] for t in tx if t["direction"] == "out")
        total_in = sum(t["amount"] for t in tx if t["direction"] == "in")
        pdf = render_simple_statement(
            title=f"Cash Flow — {month}",
            subtitle=f"{len(tx)} transactions · prepared {datetime.now().strftime('%d %b %Y')}",
            meta={"Expense": f"₹ {total_out:,.0f}", "Income": f"₹ {total_in:,.0f}",
                  "Net": f"₹ {total_in - total_out:,.0f}"},
            table_headers=["Date", "Dir", "Amount", "Head", "Company", "Mode"],
            table_rows=rows,
            footer="Generated by Mind Matters.",
        )
        fname = f"cashflow-{month}.pdf"
        caption = f"Cash flow for {month}"

    res = await tg_send_document(cid, fname, pdf, caption=caption)
    return {"ok": bool(res and res.get("ok")), "telegram": res}


# ───────────────────── Deadlines (countdown) ─────────────────────
class DeadlineIn(BaseModel):
    title: str
    due_date: str  # YYYY-MM-DD
    notes: Optional[str] = ""
    project_id: Optional[str] = None  # v2.17


class Deadline(BaseModel):
    id: str
    user_id: str
    title: str
    due_date: str
    notes: str = ""
    created_at: str


@api.get("/deadlines", response_model=List[Deadline])
async def list_deadlines(project_id: Optional[str] = None, user=Depends(get_current_user)):
    q: Dict[str, Any] = {"user_id": user["id"]}
    if project_id:
        q["project_id"] = project_id
    docs = await db.deadlines.find(q, {"_id": 0}) \
        .sort("due_date", 1).to_list(200)
    return [Deadline(**d) for d in docs]


@api.post("/deadlines", response_model=Deadline)
async def create_deadline(body: DeadlineIn, user=Depends(get_current_user)):
    project_id = body.project_id or await _ensure_default_project(user["id"])
    await _assert_project_write(user["id"], project_id)
    doc = {
        "id": new_id(), "user_id": user["id"],
        "title": body.title, "due_date": body.due_date, "notes": body.notes or "",
        "project_id": project_id,
        "created_at": now_iso(),
    }
    await db.deadlines.insert_one(dict(doc))
    return Deadline(**doc)


@api.patch("/deadlines/{did}", response_model=Deadline)
async def update_deadline(did: str, body: Dict[str, Any], user=Depends(get_current_user)):
    body = {k: v for k, v in body.items() if k in {"title", "due_date", "notes", "project_id"}}
    res = await db.deadlines.find_one_and_update(
        {"id": did, "user_id": user["id"]}, {"$set": body},
        projection={"_id": 0}, return_document=True,
    )
    if not res:
        raise HTTPException(404, "Deadline not found")
    return Deadline(**res)


@api.delete("/deadlines/{did}")
async def delete_deadline(did: str, user=Depends(get_current_user)):
    await db.deadlines.delete_one({"id": did, "user_id": user["id"]})
    return {"ok": True}


# ───────────────────── AI bulk parse (paste text or upload file) ─────────────────────
class BulkParseReq(BaseModel):
    kind: Literal["task", "routine", "expense", "loan", "investment", "note", "reminder", "deadline"]
    text: str  # could be tab-separated paste OR free text


def _parse_excel_to_rows(content: bytes) -> List[Dict[str, Any]]:
    import pandas as pd
    df = pd.read_excel(io.BytesIO(content))
    df.columns = [str(c).strip() for c in df.columns]
    return df.fillna("").to_dict(orient="records")


def _parse_csv_to_rows(content: bytes) -> List[Dict[str, Any]]:
    import csv
    f = io.StringIO(content.decode("utf-8", errors="ignore"))
    return list(csv.DictReader(f))


SCHEMA_BY_KIND = {
    "task": (
        "{date:'YYYY-MM-DD' or null (extract dates like '05/06/2026', '5 June 2026', 'tomorrow' relative to today; null only if user truly gave no date), "
        "group (free text — extract from '#<name>' (e.g. '#Work' → 'Work'), '#group:<name>', 'Group: <name>', 'group: <name>', 'in <group>'; preserve user's exact case), "
        "name (the responsible person — the one being asked TO do something; e.g. 'Rahul' from 'remind Rahul to call', 'Brinda' from 'ask Brinda for invoice'), "
        "task (a short ONE-WORD action verb in user's case e.g. 'Call', 'Send', 'Remind', 'Buy', 'Pay', 'Follow-Up'), "
        "details (the rest of the sentence — what about / what to do — KEEP EVERY non-trivial NOUN/OBJECT the user typed; e.g. 'medication for father', 'invoice', 'the bar unit'; never silently drop nouns), "
        "status:'Pending' (default), "
        "reminder_at:'YYYY-MM-DDTHH:MM' or null (ONLY when the user gave both a date AND a clock time like '4:00 pm', '16:00', '9am' — interpret 12h with am/pm; ignore timezone abbreviations like 'IST' for now since the local clock already matches user's locale)}\n"
        "EXAMPLE 1: input 'Remind rahul to send invoice #Work on 05/06/2026 4:00 pm ist' → "
        "{date:'2026-06-05', group:'Work', name:'rahul', task:'Send', details:'invoice', status:'Pending', reminder_at:'2026-06-05T16:00'}\n"
        "EXAMPLE 2: input 'call brinda about repair tomorrow' → "
        "{date:'(tomorrow's ISO date)', group:'', name:'brinda', task:'Call', details:'about repair', status:'Pending', reminder_at:null}\n"
        "EXAMPLE 3: input 'buy medication for father #Family' → "
        "{date:today, group:'Family', name:'', task:'Buy', details:'medication for father', status:'Pending', reminder_at:null}\n"
        "EXAMPLE 4: input 'pay electricity bill' → "
        "{date:null, group:'', name:'', task:'Pay', details:'electricity bill', status:'Pending', reminder_at:null}"
    ),
    "routine": (
        "{group (free text — first check for an explicit '#<name>' or '#group:<name>' hashtag and use that (e.g. '#Morning' → 'Morning'); otherwise for time-of-day words ('morning','evening','night','afternoon') USE that as group; otherwise pick a sensible category like 'Health','Work','Family'; preserve user's case), "
        "name (the person who performs it — 'Self','Wife','Father' etc. from words like 'self','me','i','wife','father','mother','children'), "
        "activity (a short verb-led title e.g. 'Walk', 'Meditate', 'Stretch'), "
        "details (location/qualifier e.g. 'at park', '30 min', 'with breathing'), "
        "frequency:'Daily'|'Weekly'|'Every 2 Weeks'|'Monthly'|'Every 2 Months'|'Quarterly'|'Half-Yearly'|'Yearly' (default 'Daily')}\n"
        "EXAMPLE 1: input 'morning walk at park daily self #Morning' → "
        "{group:'Morning', name:'Self', activity:'Walk', details:'at park', frequency:'Daily'}\n"
        "EXAMPLE 2: input 'evening meditation 20 min for wife weekly' → "
        "{group:'Evening', name:'Wife', activity:'Meditation', details:'20 min', frequency:'Weekly'}\n"
        "EXAMPLE 3: input 'father morning yoga' → "
        "{group:'Morning', name:'Father', activity:'Yoga', details:'', frequency:'Daily'}"
    ),
    "expense": (
        "{date:'YYYY-MM-DD' (default to today if missing), "
        "amount:number (handle 'lakhs'→×100000, 'crore'→×10000000, 'k'→×1000), "
        "vendor (the company/provider/merchant name e.g. 'Bajaj', 'Starbucks', 'SBI'; "
        "for insurance, this is the insurer like 'Bajaj','LIC','HDFC Ergo'), "
        "name (the person involved — beneficiary, borrower, payer; e.g. 'Karan' from 'insurance for karan'), "
        "details (free-text rest; for insurance include who it's for), "
        "head (one-word category Title-Case e.g. 'Insurance','Food','Travel','Loan','Investment','Salary'), "
        "mode ('Card'|'Cash'|'UPI'|'Bank'|'Cheque'; default 'Bank' if not specified), "
        "category ('income'|'expense'|'asset'|'liability' — use 'liability' for insurance/loans-given, "
        "'asset' for investments/loans-taken/FDs, 'income' for salary/refund/interest received, 'expense' otherwise), "
        "group (free text — extract from '#<name>' (e.g. '#Family' → 'Family') or '#group:<name>' or mention of group/category bucket; preserve user's case; else '')}\n"
        "EXAMPLE 1: 'insurance from bajaj karan 5 lakhs' → "
        "{date:today, amount:500000, vendor:'Bajaj', name:'Karan', details:'insurance for karan', head:'Insurance', mode:'Bank', category:'liability', group:''}\n"
        "EXAMPLE 2: '450 coffee at starbucks card' → "
        "{date:today, amount:450, vendor:'Starbucks', name:'', details:'coffee', head:'Food', mode:'Card', category:'expense', group:''}\n"
        "EXAMPLE 3: 'sbi fd 2 lakhs at 7.1%' → "
        "{date:today, amount:200000, vendor:'SBI', name:'', details:'FD at 7.1%', head:'Investment', mode:'Bank', category:'asset', group:''}\n"
        "EXAMPLE 4: 'lent brinda 50000 at 9%' → "
        "{date:today, amount:50000, vendor:'', name:'Brinda', details:'lent at 9%', head:'Loan', mode:'Bank', category:'asset', group:''}"
    ),
    "loan": (
        "{date:'YYYY-MM-DD', name (Title-Case), amount:number, "
        "interest:number, interest_type:'percent'|'fixed' "
        "(use 'fixed' if user wrote a flat amount like '15k', '1500', '50000 fixed' — "
        "use 'percent' if user wrote a rate like '9%', '12 percent p.a.', or no unit), "
        "reason (Title-Case), status:'Given'|'Taken'|'Pending'|'Closed', repayment_date}"
    ),
    "investment": (
        "{kind:'investment'|'insurance' (use 'insurance' if user mentioned LIC/term/health/ULIP/policy/cover, "
        "else 'investment'), "
        "type (Title-Case e.g. 'Equity','FD','Insurance','MF','Bond'), "
        "provider (Title-Case), amount_invested:number, start_date, maturity_date, "
        "rate_or_value (a string like '8% p.a.' or '₹2,40,000'), "
        "insured_for (Title-Case e.g. 'Self','Wife','Mother','Father','Medical' — only when kind='insurance'; else null), "
        "notes}"
    ),
    "note": "{title, body, tags:[lowercase] (extract from '#<tag>' hashtags like '#Personal' → 'personal'), list_title?:string, list_tag?:string, items?:[string]}. "
            "When the user wants to append bullets to an existing list (e.g. 'add milk, eggs to shopping list', "
            "'add X to <list>', or single-item buy/get/pick-up phrases like 'buy soap', 'get bread', "
            "'pick up dry cleaning', 'order coffee filters'), ALWAYS prefer list-append: set "
            "list_title='Shopping List' (or the named list if specified) AND put bullets in items[]. "
            "Recognize tag context: 'add X to #shopping' → list_tag='shopping', items=['X']. "
            "Only when user is clearly writing a fresh thought/idea/journal entry "
            "(e.g. 'idea: sunday review', 'note about meeting today') should you produce a "
            "standalone {title, body} note WITHOUT items[].",
    "reminder": (
        "{title (verb-led action, e.g. 'Call Brinda' not just 'Brinda' — interpret slang: "
        "'hit X' = 'Call X', 'ping X' = 'Message X'), "
        "fire_at_local:'YYYY-MM-DDTHH:MM' (interpret 'tomorrow 9am', '3pm', 'monday' etc. "
        "relative to today; default to 09:00 if time absent; for vague 'evening' use 18:00, "
        "'morning' 09:00, 'noon' 12:00, 'night' 21:00), "
        "notes (any extra context beyond the action — venue, why, who — keep brief), "
        "recurrence:'none'|'daily'|'weekly'|'monthly'|'quarterly'|'half-yearly'|'yearly'}"
    ),
    "deadline": "{title (Title-Case), due_date:'YYYY-MM-DD', notes}",
}


def _title_case_smart(s: str) -> str:
    # v2.1: user requested NO auto-capitalisation — feed as the user typed.
    return s if isinstance(s, str) else s


def _normalize_row(kind: str, row: Dict[str, Any]) -> Dict[str, Any]:
    """Post-process AI-parsed rows. v2.1: no casing, only trim + default-today."""
    if not isinstance(row, dict):
        return row
    for k, v in list(row.items()):
        if isinstance(v, str):
            row[k] = v.strip()
    # default date = today for tasks/expenses if missing
    if kind in ("task", "expense", "deadline") and not row.get("date"):
        row["date"] = today_key()
    return row


async def _ai_parse_bulk(text: str, kind: str) -> List[Dict[str, Any]]:
    if not EMERGENT_LLM_KEY:
        return []
    schema = SCHEMA_BY_KIND.get(kind, "{}")
    today = today_key()
    system = (
        "You are a data normalizer. Extract a JSON ARRAY of objects from the user's input. "
        f"TODAY'S DATE is {today} — use this for any 'today', 'tomorrow', 'yesterday', "
        f"'next week' references. If no date is given, leave it null (do NOT guess).\n"
        f"Each object follows this schema for kind={kind}: {schema}\n"
        "If input has multiple lines, treat each as one record. "
        "Return ONLY a JSON array, no prose."
    )
    # Chunk by line count so large pastes (50-100+ rows) don't exceed the LLM's
    # output-token budget and silently truncate (was capping at ~25 rows).
    lines = [ln for ln in text.split("\n") if ln.strip()]
    if len(lines) <= 30:
        chunks = [text]
    else:
        chunks = []
        for i in range(0, len(lines), 30):
            chunks.append("\n".join(lines[i : i + 30]))
    out: List[Dict[str, Any]] = []
    for idx, chunk in enumerate(chunks):
        chat = LlmChat(api_key=EMERGENT_LLM_KEY, session_id=f"bp-{new_id()}-{idx}",
                       system_message=system).with_model("gemini", "gemini-3-flash-preview")
        try:
            resp = await chat.send_message(UserMessage(text=chunk[:8000]))
        except Exception:
            continue
        t = resp.strip()
        if t.startswith("```"):
            t = t.strip("`")
            if t.startswith("json"):
                t = t[4:]
        try:
            arr = json.loads(t)
            if isinstance(arr, list):
                out.extend(_normalize_row(kind, r) for r in arr)
        except Exception:
            continue
    return out


@api.post("/parse/bulk")
async def parse_bulk(body: BulkParseReq, user=Depends(get_current_user)):
    """Parse pasted text into a JSON array for the given kind. UI confirms before insert."""
    rows = await _ai_parse_bulk(body.text, body.kind)
    return {"rows": rows}


@api.post("/parse/bulk-file")
async def parse_bulk_file(
    kind: str = Form(...),
    file: UploadFile = File(...),
    user=Depends(get_current_user),
):
    content = await file.read()
    name = (file.filename or "").lower()
    raw: List[Dict[str, Any]] = []
    if name.endswith((".xlsx", ".xls")):
        raw = _parse_excel_to_rows(content)
    elif name.endswith(".csv"):
        raw = _parse_csv_to_rows(content)
    else:
        raise HTTPException(400, "Upload .xlsx, .xls or .csv")
    if not EMERGENT_LLM_KEY:
        return {"rows": [], "raw_count": len(raw)}
    schema = SCHEMA_BY_KIND.get(kind, "{}")
    system = (
        "Normalize raw spreadsheet rows to a JSON array per the schema. "
        f"Schema for kind={kind}: {schema}\nReturn ONLY JSON array."
    )
    # Process in chunks of 40 rows so the LLM output token budget never clips
    # large pastes/uploads (the previous 30-row sample silently dropped rows).
    CHUNK = 40
    out_rows: List[Dict[str, Any]] = []
    for i in range(0, len(raw), CHUNK):
        chunk = raw[i : i + CHUNK]
        sample = json.dumps(chunk, default=str)
        chat = LlmChat(api_key=EMERGENT_LLM_KEY, session_id=f"bf-{new_id()}-{i}",
                       system_message=system).with_model("gemini", "gemini-3-flash-preview")
        try:
            resp = await chat.send_message(UserMessage(text=sample))
        except Exception:
            continue
        t = resp.strip()
        if t.startswith("```"):
            t = t.strip("`")
            if t.startswith("json"):
                t = t[4:]
        try:
            rows = json.loads(t)
            if isinstance(rows, list):
                out_rows.extend(_normalize_row(kind, r) for r in rows)
        except Exception:
            pass
    return {"rows": out_rows, "raw_count": len(raw)}


# ───────────────────── Invoice AI fill ─────────────────────
class InvoiceParseReq(BaseModel):
    template_id: str
    text: str


@api.post("/parse/invoice")
async def parse_invoice(body: InvoiceParseReq, user=Depends(get_current_user)):
    """Free-text → structured field-dict matching the template's required_fields."""
    if not EMERGENT_LLM_KEY:
        raise HTTPException(500, "LLM not configured")
    # find template (built-in or user)
    tpl = None
    if body.template_id in ("rkm_donation_receipt", "krm_huf_invoice"):
        tpl = next((t for t in BUILTIN_TEMPLATES if t["id"] == body.template_id), None)
    else:
        tpl = await db.templates.find_one(
            {"id": body.template_id, "user_id": user["id"]}, {"_id": 0}
        )
    if not tpl:
        raise HTTPException(404, "Template not found")
    fields = tpl.get("required_fields") or []
    field_keys = [f.get("key") for f in fields if f.get("key")]

    schema = ", ".join(f"'{k}'" for k in field_keys) or "(no fields)"
    sys_msg = (
        "Extract invoice / receipt fields from the user's input. "
        f"Return ONLY a JSON object whose keys are a subset of: {schema}. "
        "If a value is unknown, omit the key. For amounts, return numbers only. "
        "For dates, use ISO 'YYYY-MM-DD'. For line_items, return an array of "
        "{description, hsn, quantity, unit_price}. Numbers may be in words "
        "(e.g. 'fifty thousand') — convert. Capitalize proper nouns."
    )
    chat = LlmChat(api_key=EMERGENT_LLM_KEY, session_id=f"inv-{new_id()}",
                   system_message=sys_msg).with_model("gemini", "gemini-3-flash-preview")
    resp = await chat.send_message(UserMessage(text=body.text[:8000]))
    t = resp.strip()
    if t.startswith("```"):
        t = t.strip("`")
        if t.startswith("json"):
            t = t[4:]
    try:
        data = json.loads(t)
        if not isinstance(data, dict):
            data = {}
    except Exception:
        data = {}
    # filter to known keys only
    out = {k: v for k, v in data.items() if k in field_keys}
    return {"data": out, "raw": resp}


# ───────────────────── Notes — image attachments (base64) ─────────────────────
class NoteImage(BaseModel):
    id: str
    note_id: str
    user_id: str
    name: str
    mime: str
    data_b64: str
    created_at: str


@api.post("/notes/{nid}/images")
async def upload_note_image(
    nid: str,
    file: UploadFile = File(...),
    user=Depends(get_current_user),
):
    note = await db.notes.find_one({"id": nid, "user_id": user["id"]}, {"_id": 0})
    if not note:
        raise HTTPException(404, "Note not found")
    content = await file.read()
    if len(content) > 4 * 1024 * 1024:
        raise HTTPException(400, "Image too large (max 4 MB)")
    import base64
    doc = {
        "id": new_id(), "note_id": nid, "user_id": user["id"],
        "name": file.filename or "image", "mime": file.content_type or "image/png",
        "data_b64": base64.b64encode(content).decode("ascii"),
        "created_at": now_iso(),
    }
    await db.note_images.insert_one(dict(doc))
    return {k: v for k, v in doc.items() if k != "data_b64"} | {
        "data_url": f"data:{doc['mime']};base64,{doc['data_b64']}"
    }


@api.get("/notes/{nid}/images")
async def list_note_images(nid: str, user=Depends(get_current_user)):
    imgs = await db.note_images.find(
        {"note_id": nid, "user_id": user["id"]}, {"_id": 0}
    ).sort("created_at", 1).to_list(50)
    return [{**i, "data_url": f"data:{i['mime']};base64,{i['data_b64']}"} for i in imgs]


@api.delete("/notes/{nid}/images/{iid}")
async def delete_note_image(nid: str, iid: str, user=Depends(get_current_user)):
    await db.note_images.delete_one({"id": iid, "note_id": nid, "user_id": user["id"]})
    return {"ok": True}


# ───────────────────── Startup: Telegram poller + reminder loop ─────────────────────
@app.on_event("startup")
async def _mm_startup_tasks():
    # v2.0 migration: wipe loans + investments collections per user request
    try:
        dropped_l = await db.loans.delete_many({})
        dropped_i = await db.investments.delete_many({})
        if dropped_l.deleted_count or dropped_i.deleted_count:
            logger.info(f"v2 migration: wiped {dropped_l.deleted_count} loans, "
                        f"{dropped_i.deleted_count} investments")
    except Exception as e:
        logger.warning(f"startup wipe failed: {e}")

    # v2.17 migration — Item 47: Purge legacy demo data + Item 46: Ensure
    # every user has a default "Personal" project and back-fill project_id
    # on existing rows.
    try:
        await _purge_legacy_demo_data()
        await _backfill_default_projects()
    except Exception as e:
        logger.warning(f"v2.17 migration failed: {e}")

    # Indexes
    try:
        await db.users.create_index("email", unique=True, sparse=True)
        await db.projects.create_index([("owner_id", 1)])
        await db.project_members.create_index([("project_id", 1), ("user_id", 1)])
        await db.project_members.create_index("invited_email")
        await db.comments.create_index([("project_id", 1), ("resource_type", 1), ("resource_id", 1)])
    except Exception as e:
        logger.warning(f"index create: {e}")
    try:
        asyncio.create_task(tg_poll_loop(
            db, _ai_parse_text, _create_task_internal,
            _create_expense_internal, _create_note_internal,
        ))
        asyncio.create_task(reminder_loop(db))
        logger.info("Background tasks started (Telegram poller + reminders)")
    except Exception as e:
        logger.warning(f"Startup tasks failed: {e}")


# ───────────────────────────── healthcheck ─────────────────────────────
@api.get("/")
async def root():
    return {"app": "mind-matters", "ok": True, "time": now_iso()}


# register
app.include_router(api)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get("CORS_ORIGINS", "*").split(","),
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()
