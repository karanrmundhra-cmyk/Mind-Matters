"""Phase 2+ backend tests — Documents, Reminders, Telegram, Share."""
import os
import io
import time
import pytest
import requests

BASE_URL = os.environ["REACT_APP_BACKEND_URL"].rstrip("/")
API = f"{BASE_URL}/api"


@pytest.fixture(scope="module")
def client():
    s = requests.Session()
    s.headers.update({"Content-Type": "application/json"})
    r = s.post(f"{API}/auth/demo-login", json={"first_name": "Arjun"}, timeout=15)
    assert r.status_code == 200, r.text
    token = r.json()["token"]
    s.headers.update({"Authorization": f"Bearer {token}"})
    return s


# -------- News (GNews) --------
def test_news_gnews(client):
    r = client.get(f"{API}/news/headlines", timeout=15)
    assert r.status_code == 200
    body = r.json()
    assert "headlines" in body
    hl = body["headlines"]
    assert len(hl) >= 3
    # Real GNews articles should have URL field non-empty
    has_url = any((h.get("url") or "").startswith("http") for h in hl)
    assert has_url, f"Expected at least one URL in headlines: {hl}"


# -------- Documents / Templates --------
def test_list_templates_builtin(client):
    r = client.get(f"{API}/documents/templates", timeout=10)
    assert r.status_code == 200
    body = r.json()
    ids = [t["id"] for t in body["templates"]]
    assert "rkm_donation_receipt" in ids
    assert "krm_huf_invoice" in ids
    rkm = next(t for t in body["templates"] if t["id"] == "rkm_donation_receipt")
    assert isinstance(rkm.get("required_fields"), list)
    assert len(rkm["required_fields"]) > 0


def test_generate_rkm_receipt_pdf(client):
    payload = {
        "template_id": "rkm_donation_receipt",
        "data": {
            "donor_name": "TEST_Arjun Mehta",
            "donor_address": "Mumbai, India",
            "donor_pan": "ABCDE1234F",
            "sum_rupees": 5100,
            "purpose": "Annadanam",
            "receipt_no": "RKM-TEST-001",
            "receipt_date": "2026-01-15",
            "mode": "UPI",
        },
    }
    r = client.post(f"{API}/documents/generate", json=payload, timeout=30)
    assert r.status_code == 200, r.text
    assert "application/pdf" in r.headers.get("content-type", "")
    assert len(r.content) > 1024
    assert r.content[:4] == b"%PDF"


def test_generate_huf_invoice_pdf(client):
    payload = {
        "template_id": "krm_huf_invoice",
        "data": {
            "invoice_no": "INV-TEST-001",
            "invoice_date": "2026-01-15",
            "bill_to": {"name": "TEST_Client Pvt Ltd", "address": "Pune"},
            "items": [
                {"description": "Consulting", "qty": 10, "rate": 1500},
                {"description": "Setup", "qty": 1, "rate": 5000},
            ],
            "tax_percent": 18,
            "notes": "Test invoice",
        },
    }
    r = client.post(f"{API}/documents/generate", json=payload, timeout=30)
    assert r.status_code == 200, r.text
    assert "application/pdf" in r.headers.get("content-type", "")
    assert len(r.content) > 1024
    assert r.content[:4] == b"%PDF"


def test_share_telegram_without_link_400(client):
    payload = {
        "template_id": "rkm_donation_receipt",
        "data": {"donor_name": "X", "sum_rupees": 100, "receipt_no": "1",
                 "receipt_date": "2026-01-15", "mode": "Cash"},
    }
    r = client.post(f"{API}/documents/share-telegram", json=payload, timeout=15)
    assert r.status_code == 400


def test_upload_docx_template(client):
    """Build a tiny .docx with placeholders and upload it."""
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx not installed in test env")
    buf = io.BytesIO()
    d = Document()
    d.add_paragraph("Hello {{name}}, your donation of Rs {{amount}} is received.")
    d.add_paragraph("Date: {{donation_date}}")
    d.save(buf)
    buf.seek(0)
    files = {"file": ("test_tpl.docx", buf.getvalue(),
                      "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    headers = {"Authorization": client.headers["Authorization"]}
    r = requests.post(
        f"{API}/documents/templates/upload",
        files=files, data={"name": "TEST_Custom Tpl", "kind": "invoice"},
        headers=headers, timeout=30,
    )
    assert r.status_code == 200, r.text
    body = r.json()
    keys = {f["key"] for f in body.get("required_fields", [])}
    assert {"name", "amount", "donation_date"}.issubset(keys), keys
    # cleanup
    client.delete(f"{API}/documents/templates/{body['id']}")


# -------- Reminders --------
def test_reminders_crud_and_ics(client):
    fire_at = "2030-01-01T10:00:00+00:00"
    r = client.post(f"{API}/reminders",
                    json={"title": "TEST_Reminder", "notes": "do it",
                          "fire_at": fire_at, "recurrence": "none"}, timeout=10)
    assert r.status_code == 200, r.text
    rid = r.json()["id"]

    # list
    r = client.get(f"{API}/reminders", timeout=10)
    assert r.status_code == 200
    assert any(x["id"] == rid for x in r.json())

    # patch
    r = client.patch(f"{API}/reminders/{rid}", json={"title": "TEST_Reminder updated"}, timeout=10)
    assert r.status_code == 200 and r.json()["title"] == "TEST_Reminder updated"

    # ics
    r = client.get(f"{API}/reminders/{rid}/ics", timeout=10)
    assert r.status_code == 200
    body = r.text
    assert body.startswith("BEGIN:VCALENDAR")
    assert "BEGIN:VEVENT" in body and "END:VCALENDAR" in body

    # delete
    r = client.delete(f"{API}/reminders/{rid}", timeout=10)
    assert r.status_code == 200


# -------- Telegram --------
def test_tg_status(client):
    r = client.get(f"{API}/telegram/status", timeout=15)
    assert r.status_code == 200
    body = r.json()
    assert body.get("configured") is True
    assert body.get("bot_username") == "mindmattersbot", body
    # not linked initially for fresh demo user (or stays unlinked unless real flow ran)
    assert "linked" in body


def test_tg_link_code(client):
    r = client.post(f"{API}/telegram/link-code", timeout=10)
    assert r.status_code == 200
    code = r.json().get("code")
    assert isinstance(code, str) and len(code) >= 6


def test_tg_send_test_without_link_400(client):
    # Ensure unlinked first
    client.post(f"{API}/telegram/unlink", timeout=10)
    r = client.post(f"{API}/telegram/send-test", json={"text": "hi"}, timeout=10)
    assert r.status_code == 400


# -------- Share Statement --------
def test_share_statement_loan_without_link_400(client):
    client.post(f"{API}/telegram/unlink", timeout=10)
    r = client.post(f"{API}/share/statement", json={"kind": "loan"}, timeout=15)
    assert r.status_code == 400


# -------- Non-regression light checks --------
def test_dashboard_snapshot_still_works(client):
    r = client.get(f"{API}/dashboard/snapshot", timeout=15)
    assert r.status_code == 200
    for k in ("pending_tasks_count", "routine_percent_today", "insights"):
        assert k in r.json()
